#!/usr/bin/env python3
"""Generate Zamzam Museum Microclimate Control Methodology with clean (unannotated) plans."""

import sys, os
from datetime import date

_template_dir = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/_Style-Guides/Doc Style Guide"
if not os.path.exists(_template_dir):
    _template_dir = "/Users/mohamedessa/Library/Group Containers/UBF8T346G9.OneDriveStandaloneSuite/OneDrive - SAMAYA INVESTMENT.noindex/OneDrive - SAMAYA INVESTMENT/Samaya/Technical Office/_Style-Guides/Doc Style Guide"
sys.path.insert(0, _template_dir)
from samaya_doc_template import SamayaDoc

from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/Bim Unit/Zamzam Museum/Docs/07_Reports/07.1_Progress_Reports"
IMG_DIR = "/Users/mohamedessa/.hermes/tmp/zamzam_showcase_images"
TODAY = date.today().strftime("%d-%b-%Y")

CONTROL_CYCLE_SVG = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1100 520" font-family="Calibri, Arial, sans-serif">
  <defs>
    <marker id="arrow" viewBox="0 0 14 10" refX="14" refY="5" markerWidth="14" markerHeight="10" orient="auto">
      <path d="M0,0 L14,5 L0,10 Z" fill="#1E293B"/>
    </marker>
    <marker id="arrow-gray" viewBox="0 0 14 10" refX="14" refY="5" markerWidth="14" markerHeight="10" orient="auto">
      <path d="M0,0 L14,5 L0,10 Z" fill="#64748B"/>
    </marker>
  </defs>
  <rect x="30" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="130" y="60" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Showcase Air</text>
  <text x="130" y="82" text-anchor="middle" fill="#CBD5E1" font-size="13">Drawn from showcase</text>
  <text x="130" y="100" text-anchor="middle" fill="#CBD5E1" font-size="13">via intake vent</text>
  <path d="M230,70 L280,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>
  <rect x="290" y="20" width="200" height="100" rx="8" fill="#334155"/>
  <text x="390" y="60" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Filter</text>
  <text x="390" y="82" text-anchor="middle" fill="#CBD5E1" font-size="13">Particulate filtration</text>
  <text x="390" y="100" text-anchor="middle" fill="#CBD5E1" font-size="13">(HEPA / carbon)</text>
  <path d="M490,70 L540,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>
  <rect x="550" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="650" y="60" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Sensor</text>
  <text x="650" y="82" text-anchor="middle" fill="#CBD5E1" font-size="13">T/RH measurement</text>
  <text x="650" y="100" text-anchor="middle" fill="#CBD5E1" font-size="13">(every 5-10 min)</text>
  <path d="M750,70 L800,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>
  <rect x="810" y="20" width="260" height="100" rx="8" fill="#B01E2F"/>
  <text x="940" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Controller Decision</text>
  <text x="940" y="75" text-anchor="middle" fill="#FED7D7" font-size="12">Within range → Standby</text>
  <text x="940" y="95" text-anchor="middle" fill="#FED7D7" font-size="12">Outside range → Activate</text>
  <path d="M940,120 L940,150 L130,150 L130,190" stroke="#1E293B" stroke-width="3" fill="none" marker-end="url(#arrow)"/>
  <rect x="30" y="200" width="200" height="100" rx="8" fill="#334155"/>
  <text x="130" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Peltier Module</text>
  <text x="130" y="262" text-anchor="middle" fill="#CBD5E1" font-size="13">Heating / Cooling</text>
  <text x="130" y="280" text-anchor="middle" fill="#CBD5E1" font-size="13">(thermoelectric)</text>
  <path d="M230,250 L280,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>
  <rect x="290" y="200" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="390" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Humidifier /</text>
  <text x="390" y="260" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Dehumidifier</text>
  <text x="390" y="282" text-anchor="middle" fill="#CBD5E1" font-size="13">RH control module</text>
  <path d="M490,250 L540,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>
  <rect x="550" y="200" width="200" height="100" rx="8" fill="#334155"/>
  <text x="650" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Fan</text>
  <text x="650" y="262" text-anchor="middle" fill="#CBD5E1" font-size="13">Conditioned air</text>
  <text x="650" y="280" text-anchor="middle" fill="#CBD5E1" font-size="13">return to showcase</text>
  <path d="M750,250 L800,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>
  <rect x="810" y="200" width="260" height="100" rx="8" fill="#1E293B"/>
  <text x="940" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Return to Showcase</text>
  <text x="940" y="262" text-anchor="middle" fill="#CBD5E1" font-size="13">Conditioned air re-enters</text>
  <text x="940" y="280" text-anchor="middle" fill="#CBD5E1" font-size="13">showcase volume</text>
  <path d="M940,200 L940,170 L1050,170 L1050,70 L1070,70" stroke="#64748B" stroke-width="2" stroke-dasharray="6,4" fill="none" marker-end="url(#arrow-gray)"/>
  <text x="1050" y="160" text-anchor="middle" fill="#64748B" font-size="11">Continuous</text>
  <text x="1050" y="175" text-anchor="middle" fill="#64748B" font-size="11">feedback loop</text>
</svg>"""

INSTALL_SVG = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1100 520" font-family="Calibri, Arial, sans-serif">
  <defs>
    <marker id="a2" viewBox="0 0 14 10" refX="14" refY="5" markerWidth="14" markerHeight="10" orient="auto">
      <path d="M0,0 L14,5 L0,10 Z" fill="#1E293B"/>
    </marker>
  </defs>
  <rect x="30" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="130" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 1</text>
  <text x="130" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Site Survey</text>
  <text x="130" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Showcase dimensions + location</text>
  <path d="M230,70 L280,70" stroke="#1E293B" stroke-width="3" marker-end="url(#a2)"/>
  <rect x="290" y="20" width="200" height="100" rx="8" fill="#334155"/>
  <text x="390" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 2</text>
  <text x="390" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Unit Selection</text>
  <text x="390" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Size MCU per cluster volume</text>
  <path d="M490,70 L540,70" stroke="#1E293B" stroke-width="3" marker-end="url(#a2)"/>
  <rect x="550" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="650" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 3</text>
  <text x="650" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Mount MCU + Sensor</text>
  <text x="650" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Inside showcase base or plinth</text>
  <path d="M750,70 L800,70" stroke="#1E293B" stroke-width="3" marker-end="url(#a2)"/>
  <rect x="810" y="20" width="260" height="100" rx="8" fill="#334155"/>
  <text x="940" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 4</text>
  <text x="940" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Power and Data Connection</text>
  <text x="940" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Low-voltage supply + cable route</text>
  <path d="M940,120 L940,150 L130,150 L130,190" stroke="#1E293B" stroke-width="3" fill="none" marker-end="url(#a2)"/>
  <rect x="30" y="200" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="130" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 5</text>
  <text x="130" y="258" text-anchor="middle" fill="#CBD5E1" font-size="13">Setpoint Configuration</text>
  <text x="130" y="278" text-anchor="middle" fill="#CBD5E1" font-size="12">T: 20°C / RH: 50%</text>
  <path d="M230,250 L280,250" stroke="#1E293B" stroke-width="3" marker-end="url(#a2)"/>
  <rect x="290" y="200" width="200" height="100" rx="8" fill="#B01E2F"/>
  <text x="390" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 6</text>
  <text x="390" y="258" text-anchor="middle" fill="#FED7D7" font-size="13">Stabilization Test</text>
  <text x="390" y="278" text-anchor="middle" fill="#FED7D7" font-size="12">48-72h T/RH monitoring</text>
  <path d="M490,250 L540,250" stroke="#1E293B" stroke-width="3" marker-end="url(#a2)"/>
  <rect x="550" y="200" width="200" height="100" rx="8" fill="#334155"/>
  <text x="650" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 7</text>
  <text x="650" y="258" text-anchor="middle" fill="#CBD5E1" font-size="13">Calibration Verify</text>
  <text x="650" y="278" text-anchor="middle" fill="#CBD5E1" font-size="12">Compare with reference sensor</text>
  <path d="M750,250 L800,250" stroke="#1E293B" stroke-width="3" marker-end="url(#a2)"/>
  <rect x="810" y="200" width="260" height="100" rx="8" fill="#1E293B"/>
  <text x="940" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 8</text>
  <text x="940" y="258" text-anchor="middle" fill="#CBD5E1" font-size="13">Artifact Placement</text>
  <text x="940" y="278" text-anchor="middle" fill="#CBD5E1" font-size="12">Only after stable conditions</text>
</svg>"""


def add_svg_to_doc(doc, svg_content, width_cm=16.5):
    import cairosvg, tempfile
    png_data = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'), output_width=1740)
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
        f.write(png_data); t = f.name
    p = doc.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(t, width=Cm(width_cm))
    os.unlink(t)


def add_image(doc, path, width_cm=16.0, caption=""):
    if not os.path.exists(path): return
    p = doc.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(path, width=Cm(width_cm))
    if caption:
        c = doc.doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B); r.font.italic = True


def generate():
    doc = SamayaDoc()
    doc.create_header("Zamzam Museum — Showcase Microclimate Control", "ZAM-NWC-SAM-SH-XXX", "RPT", "A", TODAY)
    doc.create_footer("ZAM-NWC-SAM-SH-XXX")
    doc.add_h1("SHOWCASE MICROCLIMATE CONTROL SYSTEM — METHODOLOGY")

    doc.add_h2("1.0", "INTRODUCTION")
    doc.add_body("This document presents the recommended methodology for the microclimate control system for the Zamzam Museum showcases. The museum features 18 showcases arranged in 4 connected clusters: SHC-01 (10 units open together), SHC-02 (4 units open together), SHC-04 (2 corner units open together), and SHC-05 (2 corner units open together).")
    doc.add_body("Because showcases within each cluster share a common internal air volume, one microclimate unit per cluster is sufficient to control the entire group. Installing multiple units in the same air volume would create conflicting control loops.")

    doc.add_h2("2.0", "RECOMMENDED SOLUTION")
    doc.add_body("A dedicated microclimate control unit (MCU) installed inside or directly attached to each showcase cluster. These units provide active T/RH control with integrated monitoring.")
    doc.add_table(["Component", "Device", "Function"], [
        ["Microclimate Unit", "Showcase climate controller (e.g. CCI RH-33 or equivalent)", "Active T/RH control — heating, cooling, humidification, dehumidification"],
        ["Integrated Sensor", "Built-in T/RH sensor", "Continuous measurement, no separate logger needed"],
        ["Data Output", "USB / RS485 / optional wireless", "Data logging and export for loan compliance"],
    ], col_widths_cm=[3.0, 6.5, 7.0])

    doc.add_h2("3.0", "METHODOLOGY")
    doc.add_h3("3.1", "MICROCLIMATE CONTROL CYCLE")
    doc.add_body("Each microclimate unit operates on a closed-loop control cycle. Air is drawn from the showcase, filtered, measured by the integrated sensor, conditioned to the setpoint, and returned — maintaining a stable environment for the artifacts.")
    doc.add_body("The control cycle runs continuously at configurable intervals (default: every 5-10 minutes). If the sensor reading is within the acceptable range, the unit maintains standby mode. If outside range, the controller activates the appropriate conditioning module (Peltier for temperature, humidifier/dehumidifier for RH).")
    add_svg_to_doc(doc, CONTROL_CYCLE_SVG)

    doc.add_h3("3.2", "INSTALLATION AND COMMISSIONING METHODOLOGY")
    doc.add_body("The installation follows an 8-step methodology from site survey to artifact placement, with a total estimated duration of 5-7 days per batch.")
    doc.add_body("Each step must be completed and signed off before proceeding to the next. The stabilization test (Step 6) is the most critical — no artifacts should be placed until the logged data confirms stable T/RH within the specified range for at least 48 hours.")
    add_svg_to_doc(doc, INSTALL_SVG)

    doc.add_h2("4.0", "DISTRIBUTION PLAN — ONE MCU PER CLUSTER")
    doc.add_body("The 18 showcases are arranged in 4 connected clusters. One microclimate unit per cluster controls the entire shared air volume.")
    doc.add_table(["Cluster", "Showcases", "MCU", "Sensors", "Air Volume"],
        [
            ["SHC-01 group", "10 units (open together)", "MCU-01", "10 loggers", "Largest"],
            ["SHC-02 group", "4 units (open together)", "MCU-02", "4 loggers", "Medium"],
            ["SHC-04 corner", "2 units (open together)", "MCU-03", "2 loggers", "Small"],
            ["SHC-05 corner", "2 units (open together)", "MCU-04", "2 loggers", "Small"],
            ["Total", "18 showcases", "4 units", "18+6-8 gallery", "—"],
        ], col_widths_cm=[3.5, 4.0, 2.5, 3.0, 3.5])

    doc.add_h2("5.0", "SHOWCASE LAYOUT PLANS")
    doc.add_body("The following plans illustrate the showcase layout and arrangement across the museum galleries, showing the 4 showcase clusters (SHC-01, SHC-02, SHC-04, SHC-05) and their dimensional configuration.")

    add_image(doc, os.path.join(IMG_DIR, "page_1_resized.png"), width_cm=16.0,
              caption="Figure 1: Zamzam Museum — Showcase Layout Plan (Gallery Overview)")

    add_image(doc, os.path.join(IMG_DIR, "page_2_resized.png"), width_cm=16.0,
              caption="Figure 2: Showcase Construction Details — Glass, Frame, and Door Configuration (FR-117)")

    add_image(doc, os.path.join(IMG_DIR, "page_3_resized.png"), width_cm=16.0,
              caption="Figure 3: Showcase Cluster Plan — SHC-01, SHC-02, SHC-04, SHC-05 Arrangement and Dimensions")

    doc.add_h2("6.0", "RECOMMENDATION")
    doc.add_body("Microclimate units are recommended as the primary solution for providing active environmental control within the showcase clusters. Because the 18 showcases are arranged in 4 connected air volumes, only 4 units are required.")
    doc.add_body("Recommended next steps:")
    doc.add_body("1. Confirm exact air volume of each cluster with the showcase manufacturer")
    doc.add_body("2. Select the appropriate MCU model based on performance specifications and warranty")
    doc.add_body("3. Install and commission per the methodology in Section 3 before artifact placement")

    path = os.path.join(OUT_DIR, "Zamzam_Microclimate_Control_Methodology.docx")
    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    generate()

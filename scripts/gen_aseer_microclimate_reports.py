#!/usr/bin/env python3
"""Generate Aseer Museum Microclimate Control + T&H Monitoring Action Reports."""

import sys, os, base64
from datetime import date

# SamayaDoc template
_template_dir = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/_Style-Guides/Doc Style Guide"
if not os.path.exists(_template_dir):
    _template_dir = "/Users/mohamedessa/Library/Group Containers/UBF8T346G9.OneDriveStandaloneSuite/OneDrive - SAMAYA INVESTMENT.noindex/OneDrive - SAMAYA INVESTMENT/Samaya/Technical Office/_Style-Guides/Doc Style Guide"
sys.path.insert(0, _template_dir)
from samaya_doc_template import SamayaDoc, SamayaColors

from docx.shared import Inches, Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/Bim Unit/Aseer-Museum/04_Docs/07_Reports/07.1_Progress_Reports"

TODAY = date.today().strftime("%d-%b-%Y")

# ─── SVG Flowcharts ───────────────────────────────────────────────────────────

CONTROL_CYCLE_SVG = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1100 520" font-family="Calibri, Arial, sans-serif">
  <defs>
    <marker id="arrow" viewBox="0 0 14 10" refX="14" refY="5" markerWidth="14" markerHeight="10" orient="auto">
      <path d="M0,0 L14,5 L0,10 Z" fill="#1E293B"/>
    </marker>
    <marker id="arrow-gray" viewBox="0 0 14 10" refX="14" refY="5" markerWidth="14" markerHeight="10" orient="auto">
      <path d="M0,0 L14,5 L0,10 Z" fill="#64748B"/>
    </marker>
  </defs>
  <!-- Row 1 -->
  <rect x="30" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="130" y="60" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Showcase Air</text>
  <text x="130" y="82" text-anchor="middle" fill="#CBD5E1" font-size="13">Drawn from showcase</text>
  <text x="130" y="100" text-anchor="middle" fill="#CBD5E1" font-size="13">via intake vent</text>
  <path d="M230,70 L280,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>

  <rect x="290" y="20" width="200" height="100" rx="8" fill="#334155"/>
  <text x="390" y="60" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Filter</text>
  <text x="390" y="82" text-anchor="middle" fill="#CBD5E1" font-size="13">Particulate filtration</text>
  <text x="390" y="100" text-anchor="middle" fill="#CBD5E1" font-size="13">(HEPA / carbon)</text>
  <path d="M490,70 L540,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>

  <rect x="550" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="650" y="60" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Sensor</text>
  <text x="650" y="82" text-anchor="middle" fill="#CBD5E1" font-size="13">T/RH measurement</text>
  <text x="650" y="100" text-anchor="middle" fill="#CBD5E1" font-size="13">(every 5-10 min)</text>
  <path d="M750,70 L800,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>

  <rect x="810" y="20" width="260" height="100" rx="8" fill="#B01E2F"/>
  <text x="940" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Controller Decision</text>
  <text x="940" y="75" text-anchor="middle" fill="#FED7D7" font-size="12">Within range → Standby</text>
  <text x="940" y="95" text-anchor="middle" fill="#FED7D7" font-size="12">Outside range → Activate</text>

  <!-- Down arrow from Step 4 to Row 2 -->
  <path d="M940,120 L940,150 L130,150 L130,190" stroke="#1E293B" stroke-width="3" fill="none" marker-end="url(#arrow)"/>

  <!-- Row 2 -->
  <rect x="30" y="200" width="200" height="100" rx="8" fill="#334155"/>
  <text x="130" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Peltier Module</text>
  <text x="130" y="262" text-anchor="middle" fill="#CBD5E1" font-size="13">Heating / Cooling</text>
  <text x="130" y="280" text-anchor="middle" fill="#CBD5E1" font-size="13">(thermoelectric)</text>
  <path d="M230,250 L280,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>

  <rect x="290" y="200" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="390" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Humidifier /</text>
  <text x="390" y="260" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Dehumidifier</text>
  <text x="390" y="282" text-anchor="middle" fill="#CBD5E1" font-size="13">RH control module</text>
  <path d="M490,250 L540,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>

  <rect x="550" y="200" width="200" height="100" rx="8" fill="#334155"/>
  <text x="650" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Fan</text>
  <text x="650" y="262" text-anchor="middle" fill="#CBD5E1" font-size="13">Conditioned air</text>
  <text x="650" y="280" text-anchor="middle" fill="#CBD5E1" font-size="13">return to showcase</text>
  <path d="M750,250 L800,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow)"/>

  <rect x="810" y="200" width="260" height="100" rx="8" fill="#1E293B"/>
  <text x="940" y="240" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Return to Showcase</text>
  <text x="940" y="262" text-anchor="middle" fill="#CBD5E1" font-size="13">Conditioned air re-enters</text>
  <text x="940" y="280" text-anchor="middle" fill="#CBD5E1" font-size="13">showcase volume</text>

  <!-- Feedback loop (dashed) -->
  <path d="M940,200 L940,170 L1050,170 L1050,70 L1070,70" stroke="#64748B" stroke-width="2" stroke-dasharray="6,4" fill="none" marker-end="url(#arrow-gray)"/>
  <text x="1050" y="160" text-anchor="middle" fill="#64748B" font-size="11">Continuous</text>
  <text x="1050" y="175" text-anchor="middle" fill="#64748B" font-size="11">feedback loop</text>
</svg>"""

INSTALL_SVG = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1100 520" font-family="Calibri, Arial, sans-serif">
  <defs>
    <marker id="arrow2" viewBox="0 0 14 10" refX="14" refY="5" markerWidth="14" markerHeight="10" orient="auto">
      <path d="M0,0 L14,5 L0,10 Z" fill="#1E293B"/>
    </marker>
  </defs>
  <!-- Row 1: Steps 1-4 -->
  <rect x="30" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="130" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 1</text>
  <text x="130" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Site Survey</text>
  <text x="130" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Showcase dimensions</text>
  <path d="M230,70 L280,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow2)"/>

  <rect x="290" y="20" width="200" height="100" rx="8" fill="#334155"/>
  <text x="390" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 2</text>
  <text x="390" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Unit Selection</text>
  <text x="390" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Capacity per volume</text>
  <path d="M490,70 L540,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow2)"/>

  <rect x="550" y="20" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="650" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 3</text>
  <text x="650" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Mount Unit</text>
  <text x="650" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Inside or attached</text>
  <path d="M750,70 L800,70" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow2)"/>

  <rect x="810" y="20" width="260" height="100" rx="8" fill="#334155"/>
  <text x="940" y="55" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 4</text>
  <text x="940" y="78" text-anchor="middle" fill="#CBD5E1" font-size="13">Power Connection</text>
  <text x="940" y="98" text-anchor="middle" fill="#CBD5E1" font-size="12">Low-voltage supply</text>

  <!-- Down arrow -->
  <path d="M940,120 L940,150 L130,150 L130,190" stroke="#1E293B" stroke-width="3" fill="none" marker-end="url(#arrow2)"/>

  <!-- Row 2: Steps 5-8 -->
  <rect x="30" y="200" width="200" height="100" rx="8" fill="#1E293B"/>
  <text x="130" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 5</text>
  <text x="130" y="258" text-anchor="middle" fill="#CBD5E1" font-size="13">Setpoint Config</text>
  <text x="130" y="278" text-anchor="middle" fill="#CBD5E1" font-size="12">T: 20°C / RH: 50%</text>
  <path d="M230,250 L280,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow2)"/>

  <rect x="290" y="200" width="200" height="100" rx="8" fill="#B01E2F"/>
  <text x="390" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 6</text>
  <text x="390" y="258" text-anchor="middle" fill="#FED7D7" font-size="13">Stabilization Test</text>
  <text x="390" y="278" text-anchor="middle" fill="#FED7D7" font-size="12">48-72h monitoring</text>
  <path d="M490,250 L540,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow2)"/>

  <rect x="550" y="200" width="200" height="100" rx="8" fill="#334155"/>
  <text x="650" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 7</text>
  <text x="650" y="258" text-anchor="middle" fill="#CBD5E1" font-size="13">Calibration Verify</text>
  <text x="650" y="278" text-anchor="middle" fill="#CBD5E1" font-size="12">Compare with reference</text>
  <path d="M750,250 L800,250" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow2)"/>

  <rect x="810" y="200" width="260" height="100" rx="8" fill="#1E293B"/>
  <text x="940" y="235" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Step 8</text>
  <text x="940" y="258" text-anchor="middle" fill="#CBD5E1" font-size="13">Artifact Placement</text>
  <text x="940" y="278" text-anchor="middle" fill="#CBD5E1" font-size="12">After stable conditions</text>
</svg>"""

MONITORING_ARCH_SVG = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1100 400" font-family="Calibri, Arial, sans-serif">
  <defs>
    <marker id="arrow3" viewBox="0 0 14 10" refX="14" refY="5" markerWidth="14" markerHeight="10" orient="auto">
      <path d="M0,0 L14,5 L0,10 Z" fill="#1E293B"/>
    </marker>
  </defs>
  <!-- Tier 1: Field -->
  <rect x="30" y="20" width="1040" height="100" rx="8" fill="#1E293B"/>
  <text x="550" y="50" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Tier 1 — Field: Data Loggers</text>
  <text x="550" y="72" text-anchor="middle" fill="#CBD5E1" font-size="13">HOBO MX1101 inside each showcase (27 units) + MX1104 gallery sensors (6-8 zones)</text>
  <text x="550" y="92" text-anchor="middle" fill="#CBD5E1" font-size="13">Continuous T/RH measurement at 5-10 min intervals</text>

  <!-- Arrow down -->
  <path d="M550,120 L550,160" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow3)"/>

  <!-- Tier 2: Aggregation -->
  <rect x="30" y="170" width="1040" height="100" rx="8" fill="#334155"/>
  <text x="550" y="200" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Tier 2 — Aggregation: Gateway / Base Station</text>
  <text x="550" y="222" text-anchor="middle" fill="#CBD5E1" font-size="13">HOBO MX Gateway or Lufft Opus20</text>
  <text x="550" y="242" text-anchor="middle" fill="#CBD5E1" font-size="13">Wireless data collection from all loggers via Bluetooth Low Energy</text>

  <!-- Arrow down -->
  <path d="M550,270 L550,310" stroke="#1E293B" stroke-width="3" marker-end="url(#arrow3)"/>

  <!-- Tier 3: Cloud -->
  <rect x="30" y="320" width="1040" height="70" rx="8" fill="#1E293B"/>
  <text x="550" y="348" text-anchor="middle" fill="white" font-size="16" font-weight="bold">Tier 3 — Cloud: HOBOlink Platform</text>
  <text x="550" y="370" text-anchor="middle" fill="#CBD5E1" font-size="13">Dashboard, automated alerts, historical data, export for loan compliance</text>
</svg>"""


def render_svg_to_png(svg_content, width=1740):
    """Convert SVG string to PNG bytes using cairosvg."""
    import cairosvg
    png_data = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'), output_width=width)
    return png_data


def add_svg_to_doc(doc, svg_content, width_cm=16.5):
    """Render SVG to PNG and insert into document."""
    import tempfile
    png_data = render_svg_to_png(svg_content)
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
        f.write(png_data)
        temp_path = f.name
    p = doc.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(temp_path, width=Cm(width_cm))
    os.unlink(temp_path)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# REPORT 1: Showcase Microclimate Control System — Action Report
# ═══════════════════════════════════════════════════════════════════════════════

def generate_microclimate_report():
    doc = SamayaDoc()
    doc.create_header(
        project_name="Aseer Regional Museum",
        doc_ref="MOC-ASEER-SIC-1K0-XXX-001",
        doc_type="RPT",
        revision="A",
        date=TODAY
    )
    doc.create_footer("MOC-ASEER-SIC-1K0-XXX-001")

    doc.add_h1("SHOWCASE MICROCLIMATE CONTROL SYSTEM — ACTION REPORT")

    # 1.0 Introduction
    doc.add_h2("1.0", "INTRODUCTION")
    doc.add_body(
        "This report presents the recommended microclimate control system for the Aseer Regional Museum showcases. "
        "The museum features 27 individual showcases across multiple galleries, housing loaned artifacts, "
        "art commissions, and museum objects. All showcases are specified with passive climate control (Climate Control: P) "
        "and a recommended climate range of RH 40-60% per the NRS Showcase Schedule."
    )
    doc.add_body(
        "Unlike connected-cluster arrangements where multiple showcases share an internal air volume, "
        "the Aseer Museum showcases are individual units with separate internal volumes. "
        "Each showcase therefore requires its own microclimate control unit to maintain stable temperature and humidity conditions."
    )

    # 2.0 Recommended Solution
    doc.add_h2("2.0", "RECOMMENDED SOLUTION — SHOWCASE MICROCLIMATE UNITS")
    doc.add_body(
        "A dedicated microclimate control unit installed inside or directly attached to each showcase. "
        "These units provide active T/RH control with integrated monitoring, suitable for loaned artifacts requiring "
        "museum-grade environmental stability."
    )

    doc.add_table(
        ["Component", "Device", "Function"],
        [
            ["Microclimate Unit", "Showcase climate controller (e.g. CCI RH-33 or equivalent)", "Active T/RH control — heating, cooling, humidification, dehumidification"],
            ["Integrated Sensor", "Built-in T/RH sensor", "Continuous measurement, no separate logger needed"],
            ["Data Output", "USB / RS485 / optional wireless", "Data logging and export for loan compliance"],
        ],
        col_widths_cm=[3.0, 6.5, 7.0]
    )

    # 3.0 Methodology
    doc.add_h2("3.0", "METHODOLOGY")

    doc.add_h3("3.1", "MICROCLIMATE CONTROL CYCLE")
    doc.add_body(
        "Each microclimate unit operates on a closed-loop control cycle. Air is drawn from the showcase, "
        "filtered, measured by the integrated sensor, conditioned to the setpoint, and returned — "
        "maintaining a stable environment for the artifacts."
    )
    doc.add_body(
        "The control cycle runs continuously at configurable intervals (default: every 5-10 minutes). "
        "If the sensor reading is within the acceptable range, the unit maintains standby mode. "
        "If outside range, the controller activates the appropriate conditioning module (Peltier for temperature, "
        "humidifier/dehumidifier for RH)."
    )

    add_svg_to_doc(doc, CONTROL_CYCLE_SVG)

    doc.add_h3("3.2", "INSTALLATION AND COMMISSIONING METHODOLOGY")
    doc.add_body(
        "The installation follows an 8-step methodology from site survey to artifact placement, "
        "with a total estimated duration of 5-7 days per batch of showcases."
    )
    doc.add_body(
        "Each step must be completed and signed off before proceeding to the next. "
        "The stabilization test (Step 6) is the most critical — no artifacts should be placed "
        "until the logged data confirms stable T/RH within the specified range for at least 48 hours."
    )

    add_svg_to_doc(doc, INSTALL_SVG)

    # 4.0 Distribution Plan
    doc.add_h2("4.0", "DISTRIBUTION PLAN — ONE UNIT PER SHOWCASE")
    doc.add_body(
        "The 27 showcases are individual units with separate internal volumes. "
        "Each showcase requires its own microclimate unit. The distribution by gallery is as follows:"
    )

    doc.add_table(
        ["Gallery / Exhibit", "Showcase IDs", "Showcase Type", "Qty", "Units Required"],
        [
            ["Al Muftaha Documents & AV", "03.05_SC_01", "Type 3", "1", "1"],
            ["Al Qatt Display", "08.03_SC_01", "Type 2", "1", "1"],
            ["Scripts objects & projection", "11.03_SC_01 to 11.03_SC_08", "Type 6B / 6A", "8", "8"],
            ["Contemporary Art (Tom Nicholson)", "12.04_SC_01 to 12.04_SC_07", "Type 4", "7", "7"],
            ["Animal Sculptures & Interactive", "12.05_SC_01", "—", "1", "1"],
            ["Rock Art", "12.06_SC_01 to 12.06_SC_02", "Type 1", "2", "2"],
            ["Archaeology Display", "12.07_SC_01 to 12.07_SC_04", "Type 6B / 5A", "4", "4"],
            ["Contemporary Art (Qassem)", "LB3.04_SC_01 to LB3.04_SC_03", "Type 4", "3", "3"],
            ["Total", "", "", "27", "27"],
        ],
        col_widths_cm=[4.0, 4.5, 2.5, 1.5, 4.0]
    )

    # 5.0 Cost Estimate
    doc.add_h2("5.0", "COST ESTIMATE")
    doc.add_body(
        "Costs are estimates only. Final pricing requires supplier quotation. "
        "Branded museum-grade units (CCI RH-33 class) are recommended."
    )

    doc.add_table(
        ["Item", "Unit Price (USD)", "Qty", "Subtotal (USD)"],
        [
            ["Microclimate unit", "~$1,800", "27", "~$48,600"],
            ["Installation (per unit)", "~$200", "27", "~$5,400"],
            ["Commissioning + calibration", "~$1,500", "1 lot", "~$1,500"],
            ["Shipping", "~$1,500", "1 lot", "~$1,500"],
            ["Total", "", "", "~$57,000"],
        ],
        col_widths_cm=[5.5, 3.5, 2.0, 5.5]
    )
    doc.add_body(
        "Note: All prices are estimates and marked TBC. Actual pricing depends on selected brand, "
        "specifications, and Saudi distributor markup. Supplier quotation is required before procurement."
    )

    # 6.0 Comparison
    doc.add_h2("6.0", "COMPARISON WITH ALTERNATIVES")
    doc.add_table(
        ["Solution", "Cost", "Control", "Monitoring", "Loan Compliance"],
        [
            ["Silica gel only", "~$5,400", "Passive (buffering only)", "None", "No"],
            ["Data loggers only (HOBO)", "~$4,000", "None", "Yes", "Partial"],
            ["Microclimate units (branded)", "~$57,000", "Active T/RH control", "Built-in", "Yes"],
        ],
        col_widths_cm=[4.0, 3.0, 4.0, 3.0, 2.5]
    )

    # 7.0 Recommendation
    doc.add_h2("7.0", "RECOMMENDATION")
    doc.add_body(
        "Microclimate units are the only solution that provides active environmental control. "
        "For loaned artifacts, this is the industry standard."
    )
    doc.add_body(
        "Because the 27 showcases are individual units with separate internal volumes, "
        "each requires its own microclimate unit — 27 units total. "
        "This is the standard approach for museums with individually sealed display cases."
    )
    doc.add_body("Recommended next steps:")
    doc.add_body("1. Confirm exact internal air volume of each showcase type with the showcase manufacturer")
    doc.add_body("2. Request quotations from suppliers and local KSA distributors")
    doc.add_body("3. Select the appropriate unit based on budget and warranty")
    doc.add_body("4. Install and commission per the methodology in Section 3 before artifact placement")

    path = os.path.join(OUT_DIR, "Aseer_Microclimate_Control_Action_Report.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# REPORT 2: Temperature and Humidity Monitoring System — Action Report
# ═══════════════════════════════════════════════════════════════════════════════

def generate_monitoring_report():
    doc = SamayaDoc()
    doc.create_header(
        project_name="Aseer Regional Museum",
        doc_ref="MOC-ASEER-SIC-1K0-XXX-002",
        doc_type="RPT",
        revision="A",
        date=TODAY
    )
    doc.create_footer("MOC-ASEER-SIC-1K0-XXX-002")

    doc.add_h1("TEMPERATURE AND HUMIDITY MONITORING SYSTEM — ACTION REPORT")

    # 1.0 Introduction
    doc.add_h2("1.0", "INTRODUCTION")
    doc.add_body(
        "This report presents the recommended Temperature and Humidity (T&H) monitoring system "
        "for the Aseer Regional Museum galleries and showcases. The system is designed to protect "
        "the loaned artifacts by continuously monitoring environmental conditions inside display cases "
        "and gallery spaces, with automated alerts and a defined response protocol."
    )
    doc.add_body(
        "The museum features 27 showcases across multiple galleries (Al Muftaha, Al Qatt, Scripts, "
        "Contemporary Art, Animal Sculptures, Rock Art, Archaeology, Lobby 3). "
        "Each showcase requires individual monitoring due to the sensitivity of the artifacts on loan."
    )

    # 2.0 Recommended Devices
    doc.add_h2("2.0", "RECOMMENDED DEVICES")
    doc.add_body(
        "The following devices are recommended based on museum industry standards, accuracy requirements, "
        "and ease of installation within existing showcases."
    )

    doc.add_table(
        ["Type", "Device", "Key Specifications"],
        [
            ["In-showcase Data Logger", "HOBO MX1101 (Onset) or Testo 174H", "Bluetooth, 16,000 readings memory, 1yr battery, ±0.2°C / ±2% RH"],
            ["Gallery Data Logger", "HOBO MX1104 (ext probe) or Rotronic HF5", "On-screen display, external probe for large halls"],
            ["Gateway / Base Station", "HOBO MX Gateway or Lufft Opus20", "Wireless collection from all loggers, cloud upload"],
            ["Continuous Monitor (optional)", "Monnit Wireless T/H Sensor", "5-10yr battery, 300m range, direct cloud upload"],
        ],
        col_widths_cm=[3.5, 5.5, 7.5]
    )

    # 3.0 Monitoring Architecture
    doc.add_h2("3.0", "MONITORING ARCHITECTURE")
    doc.add_body(
        "The system follows a three-tier architecture: sensors inside each showcase and gallery, "
        "a central gateway for data aggregation, and a cloud dashboard for remote monitoring."
    )

    add_svg_to_doc(doc, MONITORING_ARCH_SVG)

    doc.add_table(
        ["Tier", "Component", "Function"],
        [
            ["1 — Field", "Data loggers inside each showcase + gallery sensors", "Continuous T/RH measurement at set intervals"],
            ["2 — Aggregation", "Gateway / base station", "Collects data wirelessly from all loggers via Bluetooth or RF"],
            ["3 — Cloud", "HOBOlink or equivalent cloud platform", "Dashboard, alerts, historical data, export"],
        ],
        col_widths_cm=[3.0, 6.0, 7.5]
    )

    # 4.0 Reading Frequency
    doc.add_h2("4.0", "READING FREQUENCY")
    doc.add_body("Reading intervals are set per location type based on artifact sensitivity and environmental stability.")

    doc.add_table(
        ["Location", "Interval", "Rationale"],
        [
            ["Inside showcases", "Every 5-10 minutes", "Artifact sensitivity, rapid detection of change"],
            ["Gallery halls", "Every 15-30 minutes", "Relatively stable, general monitoring"],
            ["Control / storage rooms", "Every 30-60 minutes", "Stable environment, low risk"],
        ],
        col_widths_cm=[4.0, 4.5, 8.0]
    )

    # 5.0 Alert Thresholds
    doc.add_h2("5.0", "ALERT THRESHOLDS AND RESPONSE PROTOCOL")

    doc.add_h3("5.1", "THRESHOLD VALUES")
    doc.add_body(
        "The following thresholds are set based on international museum standards for mixed organic and inorganic artifacts."
    )

    doc.add_table(
        ["Parameter", "Recommended Range", "Warning (Amber)", "Critical (Red)"],
        [
            ["Temperature", "18-22°C", ">24°C or <16°C", ">26°C or <14°C"],
            ["Relative Humidity", "45-55% (organic materials)", ">60% or <40%", ">65% or <35%"],
            ["T change per hour", "±2°C", "±3°C", "±5°C"],
            ["RH change per hour", "±5%", "±8%", "±10%"],
        ],
        col_widths_cm=[4.0, 4.5, 4.0, 4.0]
    )

    doc.add_h3("5.2", "RESPONSE ESCALATION")
    doc.add_table(
        ["Level", "Notification", "Action Required"],
        [
            ["Amber (Warning)", "Email + Dashboard notification", "1. Check HVAC in affected zone\n2. Verify showcase seal (air leak)\n3. Re-calibrate within 24 hours"],
            ["Red (Critical)", "Immediate SMS + Audible alarm", "1. Evacuate artifacts from affected showcase to buffer room\n2. Shut down HVAC in zone and inspect\n3. Log incident in Event Log\n4. Conservator inspects artifacts"],
            ["Power Outage", "UPS backup", "UPS powers loggers for 72+ hours continuous recording"],
        ],
        col_widths_cm=[3.0, 4.5, 9.0]
    )

    # 6.0 Logger Distribution Plan
    doc.add_h2("6.0", "LOGGER DISTRIBUTION PLAN")
    doc.add_body("Based on the approved NRS Showcase Schedule, the following distribution is proposed.")

    doc.add_table(
        ["Showcase / Area", "Showcase Type", "Qty", "Loggers Required", "Notes"],
        [
            ["03.05_SC_01 — Al Muftaha Documents", "Type 3", "1", "1", "Documents & AV display"],
            ["08.03_SC_01 — Al Qatt Display", "Type 2", "1", "1", "Al Qatt art display"],
            ["11.03_SC_01 to 11.03_SC_08 — Scripts", "Type 6B / 6A", "8", "8", "Scripts objects gallery"],
            ["12.04_SC_01 to 12.04_SC_07 — Tom Nicholson", "Type 4", "7", "7", "Contemporary art commission"],
            ["12.05_SC_01 — Animal Sculptures", "—", "1", "1", "Interactive display"],
            ["12.06_SC_01 to 12.06_SC_02 — Rock Art", "Type 1", "2", "2", "Rock art exhibits"],
            ["12.07_SC_01 to 12.07_SC_04 — Archaeology", "Type 6B / 5A", "4", "4", "Archaeology display"],
            ["LB3.04_SC_01 to LB3.04_SC_03 — Qassem", "Type 4", "3", "3", "Lobby 3 art commission"],
            ["Gallery halls", "—", "6-8 zones", "6-8", "Area sensors per gallery zone"],
            ["Total", "", "27+8 zones", "35-37 devices", "Including gateway + spares"],
        ],
        col_widths_cm=[4.5, 2.5, 1.5, 2.5, 5.5]
    )

    # 7.0 Reporting
    doc.add_h2("7.0", "REPORTING AND DOCUMENTATION")
    doc.add_table(
        ["Report Type", "Frequency", "Content"],
        [
            ["Daily Summary", "Daily", "Average T/RH per showcase + any threshold breaches"],
            ["Trend Analysis", "Weekly", "Trend analysis + week-over-week comparison"],
            ["Monthly Report", "Monthly", "Full statistics + maintenance recommendations"],
            ["Event Log", "Per incident", "Date/time/action/responsible for every alert"],
        ],
        col_widths_cm=[3.5, 3.0, 10.0]
    )

    # 8.0 Recommendation
    doc.add_h2("8.0", "RECOMMENDATION")
    doc.add_body(
        "The HOBO MX1101 system (in-showcase) with MX1104 gallery sensors and MX Gateway is recommended "
        "for the following reasons:"
    )
    doc.add_body("• Most widely used data logging system in museums globally (Louvre, British Museum, Guggenheim)")
    doc.add_body("• No wiring required inside showcases — Bluetooth Low Energy communication")
    doc.add_body("• One-year replaceable battery per logger")
    doc.add_body("• HOBOlink cloud platform with free tier for remote monitoring and alerts")
    doc.add_body("• Simple 3M adhesive mount inside showcase, positioned away from artifacts")
    doc.add_body(
        "• Estimated total cost: $5,000-$7,000 for 35-37 devices plus gateway "
        "(depending on Saudi distributor)"
    )
    doc.add_body(
        "It is recommended to procure and install the system prior to artifact placement "
        "to establish baseline environmental data and verify all alert thresholds."
    )

    path = os.path.join(OUT_DIR, "Aseer_T_H_Monitoring_Action_Report.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    p1 = generate_microclimate_report()
    p2 = generate_monitoring_report()
    print(f"\nDone. 2 reports generated:")
    print(f"  1. {p1}")
    print(f"  2. {p2}")

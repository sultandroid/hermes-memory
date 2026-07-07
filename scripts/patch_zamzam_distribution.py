#!/usr/bin/env python3
"""Update Distribution table (Sec 4) and Legend (Sec 6) to match user's manual annotations."""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

path = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/Bim Unit/Zamzam Museum/Docs/07_Reports/07.1_Progress_Reports/Zamzam_Microclimate_Control_Methodology.docx"

doc = Document(path)
body = doc.element.body

def insert_after(ref, el):
    parent = ref.getparent()
    parent.insert(list(parent).index(ref) + 1, el)

def mkp(text, size=Pt(11), bold=False, color=None, italic=False, space_after=Pt(6)):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:after'), str(int(space_after.pt * 20)))
    sp.set(qn('w:line'), '240')
    sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rf)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size.pt * 2)))
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    if italic:
        i = OxmlElement('w:i'); rPr.append(i)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
        rPr.append(c)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p

# ──────────────────────────────────────────────────────────────────
# STEP 1: Update Distribution Plan table (Section 4.0)
# ──────────────────────────────────────────────────────────────────

# Find the distribution table (table with "SHC-01 group" in first column)
for t in doc.tables:
    first_cell = t.rows[0].cells[0].text.strip() if t.rows else ''
    if 'CLUSTER' in first_cell or 'Cluster' in first_cell:
        dist_table = t
        break
else:
    print("ERROR: Distribution table not found")
    sys.exit(1)

print(f"Found distribution table with {len(dist_table.rows)} rows")

# Clear all data rows and rebuild
# Keep header, remove rows 1 to end
while len(dist_table.rows) > 1:
    row = dist_table.rows[1]
    tbl = row._tr.getparent()
    tbl.remove(row._tr)

# New data matching user's manual annotations
new_rows = [
    ["SHC-01 cluster", "3 showcases", "MCU-01", "3 loggers", "—"],
    ["SHC-02 cluster", "2 showcases", "MCU-02", "2 loggers", "—"],
    ["SHC-04 corner cluster", "4 showcases", "MCU-03", "4 loggers", "—"],
    ["SHC-05 corner cluster", "2 showcases", "MCU-04", "2 loggers", "—"],
    ["SHC-05 corner cluster", "3 showcases", "MCU-05", "3 loggers", "—"],
    ["Total", "14 showcases", "5 units", "14+6-8 gallery", "—"],
]

for i, row_data in enumerate(new_rows):
    tr = OxmlElement('w:tr')
    for j, val in enumerate(row_data):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), str(int(3.3 * 567)))
        tcPr.append(tcW)
        # alternating shading
        if i % 2 == 0:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F1F5F9')
            shading.set(qn('w:val'), 'clear')
            tcPr.append(shading)
        tc.append(tcPr)
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '40')
        pPr.append(sp); p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), 'Calibri'); rf.set(qn('w:hAnsi'), 'Calibri')
        rPr.append(rf)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20'); rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t'); t.text = val; t.set(qn('xml:space'), 'preserve')
        r.append(t); p.append(r); tc.append(p)
        tr.append(tc)
    # Add to table
    dist_table._tbl.append(tr)

print("Distribution table updated")

# ──────────────────────────────────────────────────────────────────
# STEP 2: Rebuild Legend in Section 6.0
# ──────────────────────────────────────────────────────────────────

# Find the old legend content between intro paragraph and "Figure 1"
intro_p = None
figure1_elem = None
for p in doc.paragraphs:
    if 'illustrate the showcase layout' in p.text:
        intro_p = p
    if intro_p and 'Figure 1' in p.text:
        figure1_elem = p._p
        break

if not intro_p or not figure1_elem:
    print("ERROR: Could not find legend insertion point")
    sys.exit(1)

print(f"Intro: '{intro_p.text[:50]}...'")
print(f"Figure 1: '{figure1_elem.text[:50]}...'")

# Remove everything between intro and Figure 1
intro_elem = intro_p._p
to_remove = []
found = False
for c in body:
    if c == intro_elem:
        found = True
        continue
    if c == figure1_elem:
        break
    if found:
        to_remove.append(c)

for c in to_remove:
    try:
        body.remove(c)
    except:
        pass

print(f"Removed {len(to_remove)} old legend elements")

# Insert new legend
legend_lines = [
    mkp("Legend — Device and Sensor Markings:", size=Pt(11), bold=True, space_after=Pt(2)),
    mkp("●  MCU-01  —  Microclimate Control Unit — SHC-01 cluster (3 showcases)", size=Pt(10), space_after=Pt(2)),
    mkp("●  MCU-02  —  Microclimate Control Unit — SHC-02 cluster (2 showcases)", size=Pt(10), space_after=Pt(2)),
    mkp("●  MCU-03  —  Microclimate Control Unit — SHC-04 corner cluster (4 showcases)", size=Pt(10), space_after=Pt(2)),
    mkp("●  MCU-04  —  Microclimate Control Unit — SHC-05 corner cluster (2 showcases)", size=Pt(10), space_after=Pt(2)),
    mkp("●  MCU-05  —  Microclimate Control Unit — SHC-05 corner cluster (3 showcases)", size=Pt(10), space_after=Pt(2)),
    mkp("■  Logger  —  Data Logger / T/RH Sensor (1 per showcase, inside Cluster)", size=Pt(10), space_after=Pt(2)),
    mkp("◆  Gateway —  Gateway / Base Station (central data collection point)", size=Pt(10), space_after=Pt(4)),
    mkp("Figure 4: Legend for device and sensor locations shown on the layout plans.", size=Pt(9), italic=True, color=(0x64, 0x74, 0x8B), space_after=Pt(12)),
]

current = intro_elem
for el in legend_lines:
    body.insert(list(body).index(current) + 1, el)
    current = el

# ──────────────────────────────────────────────────────────────────
# STEP 3: Update intro text showcase count 
# ──────────────────────────────────────────────────────────────────
for p in doc.paragraphs:
    if 'features 18 showcases' in p.text:
        for run in p.runs:
            if '18' in run.text:
                run.text = run.text.replace('18', '14')
        print("Updated showcase count: 18 → 14")
        break

doc.save(path)
print("\nSaved! Legend and distribution table updated to match user's manual annotations.")

# Final verification
doc2 = Document(path)
print("\n=== Verification ===")
for p in doc2.paragraphs:
    if 'Legend' in p.text or 'MCU-0' in p.text or 'Logger' in p.text or 'Gateway' in p.text:
        print(f"  {p.text[:80]}")
    if 'features' in p.text:
        print(f"  {p.text[:80]}")

for t in doc2.tables:
    if t.rows and ('SHC' in t.rows[0].cells[0].text or 'CLUSTER' in t.rows[0].cells[0].text):
        print(f"\nDistribution table:")
        for r in t.rows:
            cells = [c.text[:25].strip() for c in r.cells]
            print(f"  {' | '.join(cells)}")

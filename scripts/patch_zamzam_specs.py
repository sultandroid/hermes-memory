#!/usr/bin/env python3
"""Insert Freeair FL-Z81 specs into existing methodology DOCX without overwriting manual edits."""

import sys, os, copy
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

path = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/Bim Unit/Zamzam Museum/Docs/07_Reports/07.1_Progress_Reports/Zamzam_Microclimate_Control_Methodology.docx"

doc = Document(path)

def make_paragraph(doc, text, font_size=Pt(11), bold=False, color=None,
                   alignment=None, space_before=Pt(0), space_after=Pt(6)):
    """Create a new paragraph element and add it to the document body."""
    p = OxmlElement('w:p')
    
    # Paragraph properties
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before.pt * 20)))
    spacing.set(qn('w:after'), str(int(space_after.pt * 20)))
    pPr.append(spacing)
    if alignment is not None:
        jc = OxmlElement('w:jc')
        align_map = {WD_ALIGN_PARAGRAPH.CENTER: 'center', WD_ALIGN_PARAGRAPH.RIGHT: 'right', WD_ALIGN_PARAGRAPH.LEFT: 'left'}
        jc.set(qn('w:val'), align_map.get(alignment, 'left'))
        pPr.append(jc)
    # Bottom border for headings
    if bold and font_size >= Pt(14):
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1E293B')
        pBdr.append(bottom)
        pPr.append(pBdr)
    p.append(pPr)

    # Run
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
        rPr.append(c)
    
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    
    return p


def insert_after(element, new_element):
    """Insert new_element after element in the XML tree."""
    parent = element.getparent()
    parent.insert(list(parent).index(element) + 1, new_element)


def make_table(doc, headers, rows, col_widths_cm=None):
    """Create a table element."""
    tbl = OxmlElement('w:tbl')
    
    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(int(16.5 * 567)))  # ~16.5cm in twips
    tblPr.append(tblW)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '9CA3AF')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Grid columns
    tblGrid = OxmlElement('w:tblGrid')
    total = 16.5
    if col_widths_cm:
        for w in col_widths_cm:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(w * 567)))
            tblGrid.append(gridCol)
    else:
        col_w = total / len(headers)
        for _ in headers:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(col_w * 567)))
            tblGrid.append(gridCol)
    tbl.append(tblGrid)

    def make_cell(text, is_header=False):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), '2500')
        tcPr.append(tcW)
        if is_header:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E293B')
            shading.set(qn('w:val'), 'clear')
            tcPr.append(shading)
        tc.append(tcPr)

        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '40')
        spacing.set(qn('w:after'), '40')
        pPr.append(spacing)
        p.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Calibri')
        rFonts.set(qn('w:hAnsi'), 'Calibri')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '19' if is_header else '20')  # 9.5pt header, 10pt body
        rPr.append(sz)
        if is_header:
            b = OxmlElement('w:b')
            rPr.append(b)
            c = OxmlElement('w:color')
            c.set(qn('w:val'), 'FFFFFF')
            rPr.append(c)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    # Header row
    tr_header = OxmlElement('w:tr')
    for h in headers:
        tr_header.append(make_cell(h, is_header=True))
    tbl.append(tr_header)

    # Data rows
    for i, row in enumerate(rows):
        tr = OxmlElement('w:tr')
        for j, val in enumerate(row):
            tc = make_cell(val, is_header=False)
            if i % 2 == 0:
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F1F5F9')
                shading.set(qn('w:val'), 'clear')
                tcPr.append(shading)
            tr.append(tc)
        tbl.append(tr)

    return tbl


# Find the heading "5.0  SHOWCASE LAYOUT PLANS" and rename to "6.0"
body = doc.element.body
to_rename = []
for elem in body:
    if elem.tag == qn('w:p'):
        texts = elem.iter(qn('w:t'))
        for t in texts:
            if '5.0' in t.text and 'SHOWCASE LAYOUT' in t.text.upper():
                t.text = t.text.replace('5.0', '6.0', 1)
                to_rename.append(('sec5', elem))
            elif '6.0' in t.text and 'RECOMMENDATION' in t.text.upper():
                t.text = t.text.replace('6.0', '7.0', 1)
                to_rename.append(('sec6', elem))

print(f"Renamed: {to_rename}")

# Find the paragraph to insert after — the one just before "6.0 SHOWCASE LAYOUT PLANS"
insert_anchor = None
for elem in body:
    if elem.tag == qn('w:p'):
        texts = [t.text for t in elem.iter(qn('w:t')) if t.text]
        full = ' '.join(texts)
        if '6.0' in full and 'SHOWCASE LAYOUT' in full.upper():
            break
    insert_anchor = elem

if insert_anchor is None:
    print("ERROR: Could not find insertion anchor")
    sys.exit(1)

print(f"Insert anchor: '{''.join(t.text or '' for t in insert_anchor.iter(qn('w:t')))[:60]}...'")

# Build new Section 5 content
new_elements = []

# 5.0 heading
h2 = make_paragraph(doc, "5.0  RECOMMENDED DEVICE — FREEAIR FL-Z81".upper(),
                    font_size=Pt(14), bold=True, color=(0x1E, 0x29, 0x3B),
                    space_before=Pt(12), space_after=Pt(6))
new_elements.append(h2)

# Body
body_p = make_paragraph(doc,
    "The Freeair FL-Z81 Active Microclimate Generator is recommended for the Zamzam Museum showcases. "
    "This unit is designed specifically for museum display cases and provides active temperature and humidity control "
    "within enclosed showcase volumes.",
    font_size=Pt(11), space_after=Pt(6))
new_elements.append(body_p)

# Specs table
tbl = make_table(doc,
    ["Parameter", "Value"],
    [
        ["Model No.", "FL-Z81"],
        ["Voltage", "220-240V / 50Hz / 1PH"],
        ["Air Circulation", "90 m³/h"],
        ["Power Input", "350 W"],
        ["Operating Temperature", "5–38°C"],
        ["Refrigerant", "R410A / R407C / R134A"],
        ["Air Inlet Diameter", "30 mm"],
        ["Air Outlet Diameter", "30 mm"],
        ["Dimensions", "500 × 480 × 250 mm"],
        ["Weight", "30 kg"],
        ["Suitable Space", "≤ 0.5–1 m³"],
    ],
    col_widths_cm=[5.0, 11.5])
new_elements.append(tbl)

# Note
note_p = make_paragraph(doc,
    "Note: Specifications per manufacturer datasheet. Final selection to be confirmed against "
    "each cluster's actual internal air volume. Supplier quotation required before procurement.",
    font_size=Pt(10), color=(0x64, 0x74, 0x8B), space_after=Pt(6))
new_elements.append(note_p)

# Insert all new elements after the anchor
insert_anchor_ref = insert_anchor
for el in new_elements:
    insert_after(insert_anchor_ref, el)
    insert_anchor_ref = el

doc.save(path)
print("Saved successfully!")

# Verify
doc2 = Document(path)
for p in doc2.paragraphs:
    t = p.text.strip()
    if t and any(f"{i}.0" in t[:5] for i in range(1,8)):
        print(f"  Section: {t[:80]}")

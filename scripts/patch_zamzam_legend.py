#!/usr/bin/env python3
"""Add legend description to Section 6.0 SHOWCASE LAYOUT PLANS in the existing DOCX."""

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/Bim Unit/Zamzam Museum/Docs/07_Reports/07.1_Progress_Reports/Zamzam_Microclimate_Control_Methodology.docx"

doc = Document(path)
body = doc.element.body

def make_paragraph(text, font_size=Pt(11), bold=False, color=None, space_before=0, space_after=Pt(6), italic=False):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(space_before * 20)))
    spacing.set(qn('w:after'), str(int(space_after.pt * 20)))
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2]))
        rPr.append(c)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p

def make_table_body(headers, rows):
    tbl = OxmlElement('w:tbl')
    
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(int(16.5 * 567)))
    tblPr.append(tblW)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '9CA3AF')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    tblGrid = OxmlElement('w:tblGrid')
    col_widths = [3.5, 4.5, 8.5]
    for w in col_widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 567)))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    def make_cell(text, is_header=False):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), '2500')
        tcPr.append(tcW)
        if is_header:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1E293B')
            shading.set(qn('w:val'), 'clear')
            tcPr.append(shading)
        tc.append(tcPr)

        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '40')
        spacing.set(qn('w:after'), '40')
        pPr.append(spacing)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Calibri')
        rFonts.set(qn('w:hAnsi'), 'Calibri')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '19' if is_header else '20')
        rPr.append(sz)
        if is_header:
            b = OxmlElement('w:b'); rPr.append(b)
            c = OxmlElement('w:color'); c.set(qn('w:val'), 'FFFFFF'); rPr.append(c)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc

    tr_h = OxmlElement('w:tr')
    for h in headers:
        tr_h.append(make_cell(h, is_header=True))
    tbl.append(tr_h)

    for i, row in enumerate(rows):
        tr = OxmlElement('w:tr')
        for j, val in enumerate(row):
            tc = make_cell(val)
            if i % 2 == 0:
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F1F5F9')
                    shading.set(qn('w:val'), 'clear')
                    tcPr.append(shading)
            tr.append(tc)
        tbl.append(tr)
    return tbl


# Find "6.0  SHOWCASE LAYOUT PLANS" heading and the figure caption paragraphs
sec6_heading = None
first_figure = None
for elem in body:
    if elem.tag == qn('w:p'):
        texts = [t.text or '' for t in elem.iter(qn('w:t'))]
        full = ''.join(texts)
        if '6.0' in full and 'SHOWCASE LAYOUT' in full.upper():
            sec6_heading = elem
        if sec6_heading and 'Figure 1' in full:
            first_figure = elem
            break

if sec6_heading is None:
    print("ERROR: Section 6.0 heading not found")
    sys.exit(1)

print(f"Found Section 6.0 heading")
print(f"First figure: {first_figure is not None}")

# Insert legend content after the body text paragraph (between intro and Figure 1)
# Find the paragraph that says "The following plans illustrate..."
intro_p = None
for elem in body:
    if elem.tag == qn('w:p'):
        texts = [t.text or '' for t in elem.iter(qn('w:t'))]
        full = ''.join(texts)
        if 'illustrate the showcase layout' in full:
            intro_p = elem
            break

anchor = intro_p if intro_p else sec6_heading
print(f"Insert anchor: '{''.join(t.text or '' for t in anchor.iter(qn('w:t')))[:60]}...'")

def insert_after(ref, el):
    parent = ref.getparent()
    parent.insert(list(parent).index(ref) + 1, el)

# Build the legend content
new_items = []

# Sub-heading
new_items.append(make_paragraph("", font_size=Pt(6)))

# Legend table
legend_headers = ["Symbol", "Marking", "Description"]
legend_rows = [
    ["●", "MCU-01", "Microclimate Control Unit — SHC-01 cluster (10 showcases)"],
    ["●", "MCU-02", "Microclimate Control Unit — SHC-02 cluster (4 showcases)"],
    ["●", "MCU-03", "Microclimate Control Unit — SHC-04 corner cluster (2 showcases)"],
    ["●", "MCU-04", "Microclimate Control Unit — SHC-05 corner cluster (2 showcases)"],
    ["■", "Logger", "Data Logger / T/RH Sensor — 1 per showcase, inside case"],
    ["◆", "Gateway", "Gateway / Base Station — central data collection point"],
]
new_items.append(make_table_body(legend_headers, legend_rows))

new_items.append(make_paragraph(
    "Figure 4: Legend for device and sensor locations shown on the layout plans.",
    font_size=Pt(9), color=(0x64, 0x74, 0x8B), italic=True, space_after=Pt(12)
))

# Insert after anchor
current = anchor
for el in new_items:
    insert_after(current, el)
    current = el

doc.save(path)
print("Saved successfully!")

# Verify
doc2 = Document(path)
fig_count = 0
for p in doc2.paragraphs:
    if 'Figure' in p.text and any(d in p.text for d in '1234'):
        fig_count += 1
        print(f"  {p.text[:80]}")
print(f"Total figures: {fig_count}")

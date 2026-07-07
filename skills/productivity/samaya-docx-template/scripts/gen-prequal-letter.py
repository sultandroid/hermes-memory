#!/usr/bin/env python3
"""
Generate a prequalification letter for a subcontractor submitting TO Samaya.
Usage: python3 gen-prequal-letter.py
Edit the constants at the top for each use case.
"""

import os, sys
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── EDIT THESE ──
COMPANY = "Landscape Evergreen"
SCOPE = "landscaping and external works"
DOC_REF = "MOC-ASEER-SIC-1K0-PQ-00XX"
PROJECT = "Aseer Regional Museum — Contract 0010003521"
DATE = "July 2026"
INCLUDE_RACI = True
INCLUDE_RISK = True
# ── END EDIT ──

NAVY = RGBColor(0x1E, 0x29, 0x3B)
DARK_GRAY = RGBColor(0x33, 0x41, 0x55)
MED_GRAY = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shading)

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.0)

def add_h1(text):
    p = doc.add_paragraph()
    p.space_before = Pt(18); p.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = NAVY; run.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="1E293B"/></w:pBdr>')
    pPr.append(pBdr)

def add_h2(number, text):
    p = doc.add_paragraph()
    p.space_before = Pt(14); p.space_after = Pt(4)
    run = p.add_run(f'{number}  {text.upper()}')
    run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = NAVY; run.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="2" w:color="334155"/></w:pBdr>')
    pPr.append(pBdr)

def add_body(text):
    p = doc.add_paragraph()
    p.space_after = Pt(6); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.color.rgb = BLACK; run.font.name = 'Calibri'

def add_signed(text):
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.color.rgb = BLACK; run.font.name = 'Calibri'

# ── Cover ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(80)
run = p.add_run("PREQUALIFICATION LETTER")
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = NAVY; run.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.space_before = Pt(6)
r2 = p2.add_run(PROJECT)
r2.font.size = Pt(14); r2.font.color.rgb = DARK_GRAY; r2.font.name = 'Calibri'

for label, val in [("Document Ref:", DOC_REF), ("Submitted to:", "Samaya Investment — Technical Office"),
                    ("Project:", PROJECT), ("Date:", DATE)]:
    pm = doc.add_paragraph(); pm.alignment = WD_ALIGN_PARAGRAPH.CENTER; pm.space_after = Pt(2)
    rl = pm.add_run(label + "  "); rl.font.size = Pt(10); rl.font.color.rgb = MED_GRAY; rl.font.name = 'Calibri'
    rv = pm.add_run(val); rv.font.size = Pt(10); rv.font.color.rgb = DARK_GRAY; rv.font.name = 'Calibri'

p5 = doc.add_paragraph(); p5.alignment = WD_ALIGN_PARAGRAPH.CENTER; p5.space_before = Pt(20)
r5 = p5.add_run(f"Prepared by:  {COMPANY}")
r5.font.size = Pt(11); r5.font.bold = True; r5.font.color.rgb = NAVY; r5.font.name = 'Calibri'

doc.add_page_break()

# ── 1. Introduction ──
add_h1("1. Introduction")
add_body(f"{COMPANY} is pleased to submit this prequalification letter for the "
         f"{PROJECT}. We confirm our interest in being considered as a qualified "
         f"subcontractor for the {SCOPE} scope of this prestigious project.")

# ── 2. Project Understanding ──
add_h1("2. Project Understanding")
add_body("We have reviewed the project scope, design intent, and technical specifications "
         "provided by Samaya Investment. We confirm our full understanding of:")
add_body("- The project's significance as a regional museum under the Aseer Development Authority.\n"
         "- The design vision led by NRS (Nissen Richards Studio) and the architectural coordination requirements.\n"
         "- The project programme and milestone schedule per the master programme.\n"
         "- The quality standards, material specifications, and finishing requirements for all works.\n"
         "- The contractual framework, including Samaya's role as main contractor and the CG supervision structure.")

# ── 3. Design Understanding ──
add_h1("3. Design Understanding")
add_body("We confirm our understanding of the project design, including:")
add_body("- The design intent, material palettes, and finishes as specified in the design package.\n"
         "- The integration of our works with architectural, MEP, and civil infrastructure elements.\n"
         "- The coordination requirements with NRS design team and Samaya's BIM unit.\n"
         "- The material submittal and approval process through the CG review cycle.\n"
         "- The requirement to produce shop drawings and method statements prior to execution.")

# ── 4. Execution Sequence ──
add_h1("4. Execution Sequence Understanding")
add_body("We confirm our understanding of the project execution sequence:")
add_body("- The phased delivery approach aligned with the master programme milestones.\n"
         "- The sequence of enabling works, service coordination, and installation.\n"
         "- The requirement to coordinate with MEP, civil, and building works contractors.\n"
         "- The inspection and testing regime at each stage, including hold points for CG approval.\n"
         "- The commissioning and handover procedures, including snagging, defect liability, and as-built documentation.")

# ── 5. Compliance ──
add_h1("5. Compliance Statement")
add_body(f"{COMPANY} hereby confirms our commitment to comply with all project requirements, including:")
add_body("- All technical specifications, material standards, and quality control procedures.\n"
         "- The project quality management system and inspection/test plans (ITP).\n"
         "- Health, safety, and environmental (HSE) requirements, including site-specific safety plans.\n"
         "- The project programme and milestone dates.\n"
         "- All submittal, approval, and documentation procedures per the project BEP.\n"
         "- The CG review and approval process for all design submittals, material samples, and method statements.\n"
         "- All applicable Saudi Building Code (SBC) and local authority requirements.")

# ── 6. Capability ──
add_h1("6. Company Capability")
add_body(f"{COMPANY} brings the following capabilities to the project:")
add_body("- Proven experience in large-scale projects in the Kingdom of Saudi Arabia.\n"
         "- Qualified engineers and site supervisors with relevant project experience.\n"
         "- Access to specialised equipment and plant.\n"
         "- Established supply chain for quality materials.\n"
         "- Commitment to quality, safety, and timely delivery.")

# ── 7. RACI Matrix ──
if INCLUDE_RACI:
    add_h1("7. RACI Matrix")
    add_body("The following RACI matrix defines roles and responsibilities. "
             "R = Responsible, A = Accountable, C = Consulted, I = Informed.")

    raci_headers = ["Activity", f"{COMPANY.split()[0]}\n{COMPANY.split()[-1]}", "NRS", "Samaya\nBIM Unit", "Samaya\nPM", "CG"]
    raci_rows = [
        ["Design development", "R", "A", "C", "C", "I"],
        ["Shop drawings & method statements", "R", "C", "A", "C", "I"],
        ["Material submittals & samples", "R", "A", "C", "C", "I"],
        ["BIM coordination (LOD 300)", "C", "C", "R/A", "I", "I"],
        ["Service coordination", "R", "C", "C", "A", "I"],
        ["MEP interface coordination", "C", "C", "R", "A", "I"],
        ["Quality control / ITP execution", "R", "I", "C", "A", "I"],
        ["HSE plan & site safety", "R", "I", "I", "A", "I"],
        ["Progress reporting", "R", "I", "I", "A", "C"],
        ["Snagging & defect rectification", "R", "C", "I", "A", "I"],
        ["As-built documentation", "R", "C", "A", "C", "I"],
        ["Commissioning & handover", "R", "C", "C", "A", "A"],
    ]

    table = doc.add_table(rows=1 + len(raci_rows), cols=len(raci_headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(raci_headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(8); run.font.bold = True; run.font.color.rgb = WHITE; run.font.name = 'Calibri'
        set_cell_shading(cell, '1E293B')
    for ri, row_data in enumerate(raci_rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]; cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8); run.font.color.rgb = BLACK; run.font.name = 'Calibri'
            if ri % 2 == 1:
                set_cell_shading(cell, 'F1F5F9')

# ── 8. Risk Register ──
if INCLUDE_RISK:
    section_num = "8" if INCLUDE_RACI else "7"
    add_h1(f"{section_num}. Risk Register")
    add_body(f"{COMPANY} has identified the following project-specific risks and proposed mitigation measures.")

    risk_headers = ["#", "Risk Description", "Likelihood", "Impact", "Severity", "Mitigation Measure", "Owner"]
    risk_rows = [
        ["R1", "Imported material lead times exceed programme", "Medium", "High", "High",
         "Place long-lead orders immediately on award; maintain buffer stock", "Procurement"],
        ["R2", "Weather window closure delays installation", "Medium", "Medium", "Medium",
         "Schedule in optimal season; have fallback scope ready", "Site Manager"],
        ["R3", "Underground services clash with works", "Medium", "High", "High",
         "Full utility survey before excavation; BIM coordination; mark routes on site", "BIM Coord."],
        ["R4", "CG rejection of material samples delays start", "Medium", "Medium", "Medium",
         "Submit 3 alternatives per material; pre-approve with NRS before CG submittal", "Design Lead"],
        ["R5", "Coordination delays with other contractors", "Low", "High", "Medium",
         "Weekly interface meetings; phased handover; clear interface register", "PM"],
        ["R6", "Site access constraints limit equipment movement", "Low", "Medium", "Low",
         "Logistics plan with designated routes; coordinate deliveries with main contractor", "Site Manager"],
        ["R7", "Design changes during construction", "Low", "Medium", "Low",
         "Freeze design at IFC gate; manage changes through formal RFI process", "Design Lead"],
        ["R8", "Labour availability during peak season", "Medium", "Medium", "Medium",
         "Maintain crew buffer (20%); cross-train teams", "HR / Admin"],
    ]

    table2 = doc.add_table(rows=1 + len(risk_rows), cols=len(risk_headers))
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(risk_headers):
        cell = table2.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(7.5); run.font.bold = True; run.font.color.rgb = WHITE; run.font.name = 'Calibri'
        set_cell_shading(cell, '1E293B')
    for ri, row_data in enumerate(risk_rows):
        for ci, val in enumerate(row_data):
            cell = table2.rows[ri + 1].cells[ci]; cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci in (1, 5) else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7.5); run.font.color.rgb = BLACK; run.font.name = 'Calibri'
            if ri % 2 == 1:
                set_cell_shading(cell, 'F1F5F9')

# ── Declaration ──
decl_num = "9" if INCLUDE_RACI and INCLUDE_RISK else "7" if not INCLUDE_RACI and not INCLUDE_RISK else "8"
add_h1(f"{decl_num}. Declaration")
add_body(f"We, {COMPANY}, declare that the information provided in this prequalification letter "
         f"is accurate and complete. We confirm our readiness to proceed with the {SCOPE} "
         f"for the {PROJECT} and look forward to the opportunity to contribute to this landmark project.")

doc.add_paragraph()
add_signed("Yours faithfully,")
doc.add_paragraph()
add_signed("____________________________")
add_signed("Authorised Signatory")
add_signed(COMPANY)
add_signed("Date: ________ / ________ / 2026")

# ── Save ──
out_path = f"/tmp/{COMPANY.replace(' ', '_')}_Prequalification_Letter.docx"
doc.save(out_path)
print(f"Saved to {out_path}")

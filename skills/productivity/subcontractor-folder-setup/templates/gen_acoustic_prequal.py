#!/usr/bin/env python3
"""
Generate Acoustic Specialist Prequalification Support Documents
for multiple companies — Samaya DOCX style with SOW + RACI matrix.

Usage:
    python3 gen_acoustic_prequal.py

Output:
    00_Prequalification/<Company>_Prequalification_Support_Document.docx
    for each company in COMPANIES list.

Dependencies:
    pip install python-docx
    SamayaDoc template at:
    ~/.../Technical Office/_Style-Guides/Doc Style Guide/samaya_doc_template.py
"""
import sys, os

_template_dir = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/_Style-Guides/Doc Style Guide"
if not os.path.exists(_template_dir):
    _template_dir = "/Users/mohamedessa/Library/Group Containers/UBF8T346G9.OneDriveStandaloneSuite/OneDrive - SAMAYA INVESTMENT.noindex/OneDrive - SAMAYA INVESTMENT/Samaya/Technical Office/_Style-Guides/Doc Style Guide"
sys.path.insert(0, _template_dir)
from samaya_doc_template import SamayaDoc, SamayaColors, format_header_cell, format_data_cell, set_cell_shading
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── CONFIG ────────────────────────────────────────────────────
BASE = "/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/Bim Unit/Aseer-Museum/24_Subcontractors/18_Acoustic_Specialist"

COMPANIES = [
    {"name": "ACOUSTIEG", "folder": "10_ACOUSTIEG", "pq_ref": "MOC-MUS-ASE-1A0-PQ-0123", "doc_ref": "MOC-ASEER-SIC-1K0-SC-019-PREQ-001"},
    {"name": "AME", "folder": "11_AME", "pq_ref": "MOC-MUS-ASE-1A0-PQ-0124", "doc_ref": "MOC-ASEER-SIC-1K0-SC-019-PREQ-002"},
    {"name": "JOCAVI", "folder": "12_JOCAVI", "pq_ref": "MOC-MUS-ASE-1A0-PQ-0125", "doc_ref": "MOC-ASEER-SIC-1K0-SC-019-PREQ-003"},
]

# ── RACI DATA ────────────────────────────────────────────────
RACI_ACTIVITIES = [
    "Acoustic Assessment & Survey", "Acoustic Design Development",
    "Material Selection & Specification", "BIM Coordination (LOD 300)",
    "Submittal Preparation & Review", "CG Review & Approval",
    "Material Procurement & Samples", "Site Installation Witnessing",
    "Inspection & Testing (ISO 3382-1)", "Commissioning & Certification",
    "Handover Documentation & O&M",
]

RACI_MATRIX = {
    "Acoustic Assessment & Survey":           {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"I","MEP Coord":"C","BIM Unit":"I"},
    "Acoustic Design Development":            {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"C","MEP Coord":"C","BIM Unit":"I"},
    "Material Selection & Specification":     {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"C","MEP Coord":"I","BIM Unit":"I"},
    "BIM Coordination (LOD 300)":             {"Specialist":"C","Samaya PM":"I","NRS":"C","CG":"I","MEP Coord":"C","BIM Unit":"R/A"},
    "Submittal Preparation & Review":         {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"C","MEP Coord":"I","BIM Unit":"C"},
    "CG Review & Approval":                   {"Specialist":"C","Samaya PM":"C","NRS":"C","CG":"A","MEP Coord":"I","BIM Unit":"I"},
    "Material Procurement & Samples":         {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"C","MEP Coord":"I","BIM Unit":"I"},
    "Site Installation Witnessing":           {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"C","MEP Coord":"C","BIM Unit":"I"},
    "Inspection & Testing (ISO 3382-1)":      {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"C","MEP Coord":"I","BIM Unit":"I"},
    "Commissioning & Certification":          {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"A","MEP Coord":"I","BIM Unit":"I"},
    "Handover Documentation & O&M":           {"Specialist":"R","Samaya PM":"A","NRS":"C","CG":"C","MEP Coord":"I","BIM Unit":"I"},
}

RACI_ROLES = ["Specialist", "Samaya PM", "NRS", "CG", "MEP Coord", "BIM Unit"]

# ── HELPERS ──────────────────────────────────────────────────

def add_raci_table(doc, num, title):
    doc.add_h2(num, title)
    doc.add_body("R = Responsible | A = Accountable | C = Consulted | I = Informed")
    cols = 1 + len(RACI_ROLES)
    rows = 1 + len(RACI_ACTIVITIES)
    table = doc.doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(["Activity"] + RACI_ROLES):
        format_header_cell(table.cell(0, i), h)
    for r_idx, activity in enumerate(RACI_ACTIVITIES):
        format_data_cell(table.cell(r_idx + 1, 0), activity, WD_ALIGN_PARAGRAPH.LEFT)
        row_data = RACI_MATRIX[activity]
        for c_idx, role in enumerate(RACI_ROLES):
            format_data_cell(table.cell(r_idx + 1, c_idx + 1), row_data.get(role, "-"), WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_body("")

def make_info_table(doc, data):
    t = doc.doc.add_table(rows=len(data), cols=2)
    for i, (k, v) in enumerate(data):
        format_data_cell(t.cell(i, 0), k, WD_ALIGN_PARAGRAPH.LEFT)
        format_data_cell(t.cell(i, 1), v, WD_ALIGN_PARAGRAPH.LEFT)
        if i == 0:
            set_cell_shading(t.cell(i, 0), "1E293B")
            set_cell_shading(t.cell(i, 1), "1E293B")
            for p in [t.cell(i,0).paragraphs[0], t.cell(i,1).paragraphs[0]]:
                for r in p.runs:
                    r.font.color.rgb = SamayaColors.WHITE
                    r.font.bold = True
    doc.add_body("")

def make_2col_table(doc, headers, rows_data):
    t = doc.doc.add_table(rows=1+len(rows_data), cols=2)
    format_header_cell(t.cell(0, 0), headers[0])
    format_header_cell(t.cell(0, 1), headers[1])
    for i, (c1, c2) in enumerate(rows_data):
        format_data_cell(t.cell(i+1, 0), c1, WD_ALIGN_PARAGRAPH.LEFT)
        format_data_cell(t.cell(i+1, 1), c2, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_body("")

def make_3col_table(doc, headers, rows_data):
    t = doc.doc.add_table(rows=1+len(rows_data), cols=3)
    for i, h in enumerate(headers):
        format_header_cell(t.cell(0, i), h)
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            format_data_cell(t.cell(i+1, j), val, align)
    doc.add_body("")

# ── GENERATOR ────────────────────────────────────────────────

def generate(company):
    name = company["name"]
    folder = company["folder"]
    pq_ref = company["pq_ref"]
    doc_ref = company["doc_ref"]

    out_dir = os.path.join(BASE, folder, "00_Prequalification")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{name}_Prequalification_Support_Document.docx")

    doc = SamayaDoc()
    doc.create_header("Aseer Regional Museum", doc_ref, "PREQ", "R00", "July 2026")
    doc.create_footer(doc_number=doc_ref, confidential=True)

    doc.add_h1("PREQUALIFICATION SUPPORT DOCUMENT")
    doc.add_h2_u("ACOUSTIC SPECIALIST CONSULTANT")
    doc.add_body("")

    make_info_table(doc, [
        ("Project", "Aseer Regional Museum — Abha, Aseer Region, KSA"),
        ("Main Contractor", "Samaya Investment — Technical Office, BIM Unit"),
        ("Contract No.", "0010003521"),
        ("Trade", "Acoustic Specialist Consultant (SC-019)"),
        ("Proposed Specialist", name),
        ("Prequalification Ref.", pq_ref),
        ("Document Ref.", doc_ref),
        ("Issue Date", "July 2026"),
    ])

    doc.add_body(f"Submitted to: Samaya Investment — Technical Office, BIM Unit")
    doc.add_body(f"Proposed Subcontractor: {name}")
    doc.add_body("")

    # 1. Introduction
    doc.add_h2("1.0", "INTRODUCTION & PURPOSE")
    doc.add_body(
        f"Samaya Investment invites {name} to submit a prequalification proposal for the "
        f"Acoustic Specialist Consultant role on the Aseer Regional Museum project. "
        f"This document provides the project scope, design deliverables, RACI responsibility matrix, "
        f"and supporting reference documents required to prepare a complete submission."
    )
    doc.add_body(
        f"{name} is requested to review the scope below, complete the submission requirements, "
        f"and return this document with company stamp and authorised signature as confirmation "
        f"of capability and intent to proceed."
    )

    # 2. Project Overview
    doc.add_h2("2.0", "PROJECT OVERVIEW")
    doc.add_body(
        "The Aseer Regional Museum is a museum fit-out project located in Abha, Aseer Region, "
        "Kingdom of Saudi Arabia, at an altitude of approximately 2,200 metres. The project involves "
        "the design and construction of museum galleries, auxiliary spaces, and associated facilities "
        "under a Design & Build, Lump-Sum Milestone-Based contract."
    )
    make_info_table(doc, [
        ("Project Name", "Aseer Regional Museum of Art"),
        ("Location", "Abha, Aseer Region, KSA (Altitude ~2,200m)"),
        ("Contract Type", "Design & Build, Lump-Sum Milestone-Based"),
        ("Main Contractor", "Samaya Investment"),
        ("Employer", "Ministry of Culture (MoC), KSA"),
        ("PMC", "ACE Moharram-Bakhoum"),
        ("Design Lead / AoR", "Niels Rasmussen Studio (NRS)"),
        ("Supervising Engineer", "Consultant Group (CG)"),
        ("Contract Period", "10 months from December 2025"),
    ])

    # 3. Scope of Work
    doc.add_h2("3.0", "ACOUSTIC SCOPE OF WORK")
    doc.add_body(
        "Per the Main Contract Scope of Work (SoW Section 3.5, Section 8.1) and Employer Requirements "
        "(ER Section 3.5), the Acoustic Specialist Consultant is responsible for the following:"
    )

    doc.add_h3("3.1", "Acoustic Assessment")
    make_2col_table(doc, ["Item", "Description / Standard"], [
        ("Background noise levels", "NC/NR per space (galleries, AV spaces, cafe, library, events) — BS 8233:2014, ISO 3382-1:2009"),
        ("Reverberation time (RT60)", "Targets per space per iAcoustics Strategy V2 (ref. J3574) — DIN 18041"),
        ("Speech intelligibility (STI)", "For AV-driven spaces (events, projection rooms) — IEC 60268-16:2020"),
        ("Sound transmission", "STC/Rw between adjacent spaces — ISO 717"),
        ("HVAC noise criteria", "Duct silencer specification, low-noise diffuser/grille selection — ASHRAE, SBC 501"),
    ])

    doc.add_h3("3.2", "Design Recommendations")
    make_2col_table(doc, ["Item", "Description"], [
        ("Ceiling treatments", "Spray-applied acoustic treatment (Class B), acoustic ceiling panels (Class A), metal baffles (Class C), perforated plasterboard (Class B)"),
        ("Wall treatments", "Class A acoustic wall panels, fabric-wrapped panels, perforated panels, diffusers, absorbers"),
        ("Floor treatment", "Impact noise control"),
        ("Door/partition ratings", "STC ratings per room type"),
        ("HVAC attenuation", "Duct silencers, linings, vibration isolators"),
        ("AV coordination", "Speaker placement coordination with NRS"),
        ("Material specification", "Per NRS Acoustic Materials Spec Rev 00 — Tier A/B/C zones, NRC >= 0.85, fire EN 13501-1, GREENGUARD Gold"),
    ])

    doc.add_h3("3.3", "Commissioning Support")
    make_2col_table(doc, ["Deliverable", "Timing"], [
        ("Witness installation of acoustic treatments", "Continuous (weeks 150-210)"),
        ("Post-install acoustic measurements (ISO 3382-1)", "TOC - 30 days"),
        ("Compliance report (measured vs design targets)", "TOC - 14 days"),
        ("Acoustic commissioning certificate", "At TOC"),
    ])

    # 4. Deliverables
    doc.add_h2("4.0", "DESIGN DELIVERABLES REQUIRED")
    doc.add_body("The following deliverables are required per the project design stages:")
    make_3col_table(doc, ["Ref", "Deliverable", "Phase"], [
        ("AC-001", "Acoustic design criteria document — NR/RT/STC targets per space", "50% Design"),
        ("AC-002", "Existing acoustic conditions assessment — site measurements, baseline noise levels", "50% Design"),
        ("AC-003", "Preliminary acoustic treatment strategy — wall, ceiling, floor treatment per gallery", "50% Design"),
        ("AC-004", "Acoustic study report — reverberation time calculations, speech intelligibility", "90% Design"),
        ("AC-005", "Noise control report — MEP background noise, AV sound levels, cross-talk", "90% Design"),
        ("AC-006", "Wall acoustic treatment specification — fabric panels, perforated panels, STC ratings", "90% Design"),
        ("AC-007", "Ceiling acoustic treatment specification — acoustic baffles, cloud panels", "90% Design"),
        ("AC-008", "Acoustic treatment layout drawings — locations per gallery and auxiliary space", "90% Design"),
        ("AC-009", "Sound insulation details — partition build-up, flanking paths, seals", "90% Design"),
        ("AC-010", "Coordination with AV — acoustic requirements for projection rooms, AV spaces", "90% Design"),
        ("AC-011", "Final acoustic design package — fully coordinated with all trades", "100% Design"),
        ("AC-012", "Acoustic material schedule — products, thicknesses, NRC/STC ratings", "100% Design"),
        ("AC-013", "MEP noise and vibration control specification — isolators, silencers", "100% Design"),
        ("AC-014", "Acoustic product submittals — datasheets, test reports per product", "IFC/AFC"),
        ("AC-015", "Material samples — acoustic panels, fabric wraps, baffle finishes", "IFC/AFC"),
        ("AC-016", "ITP — Inspection and Test Plan for acoustic installation", "IFC/AFC"),
        ("AC-017", "AFC Documentation — Designer certification", "IFC/AFC"),
        ("AC-018", "Site acoustic testing — on-site verification of RT, STC, background noise", "Construction"),
        ("AC-019", "Acoustic commissioning report — measured vs design targets, certification", "Handover"),
        ("AC-020", "Record Drawings / As-Built — acoustic treatments", "Handover"),
        ("AC-021", "O&M manual — acoustic panel cleaning, replacement", "Handover"),
    ])

    # 5. RACI Matrix
    add_raci_table(doc, "5.0", "RESPONSIBILITY (RACI) MATRIX")
    doc.add_body(
        "The RACI matrix above defines roles and responsibilities for acoustic works across the project lifecycle. "
        f"{name} as the Acoustic Specialist is Responsible (R) for technical delivery, with Samaya PM "
        "Accountable (A) for overall coordination. NRS and CG are consulted/approving parties."
    )

    # 6. Prequalification Requirements
    doc.add_h2("6.0", "PREQUALIFICATION REQUIREMENTS")
    doc.add_h3("6.1", "Mandatory Certifications")
    make_2col_table(doc, ["Certification", "Requirement"], [
        ("ISO 9001:2015", "Quality management system"),
        ("ISO 14001:2015", "Environmental management"),
        ("ISO 45001:2018", "Health and safety management"),
        ("SBC Compliance", "SBC 201, SBC 301, SBC 501, SBC 801"),
        ("SCE Registration", "Saudi Council of Engineers — lead acoustician"),
    ])
    doc.add_h3("6.2", "Experience Requirements")
    make_2col_table(doc, ["Criteria", "Minimum"], [
        ("Museum/cultural projects completed", "3 projects minimum in last 10 years"),
        ("Museum acoustic design experience", "5+ years in museum/gallery acoustic design"),
        ("KSA project experience", "Preferred — 2+ projects in Kingdom"),
        ("Similar scale", "Projects >= 3,000 m2 treated area"),
        ("Commissioning experience", "Post-install acoustic measurement and certification"),
    ])
    doc.add_h3("6.3", "Technical Capabilities")
    make_2col_table(doc, ["Capability", "Required"], [
        ("Acoustic modelling software", "CadnaR, Treble (wave-based), or equivalent"),
        ("Measurement equipment", "ISO 3382-1 compliant RT measurement, sound level meters (Class 1)"),
        ("Fire testing knowledge", "EN 13501-1, ASTM E84"),
        ("Material testing", "ISO 354 / ASTM C423 sound absorption, Oddy test (conservation-grade)"),
        ("BIM capability", "Revit model review, Navisworks clash detection"),
        ("GREENGUARD Gold", "Knowledge of VOC/emissions certification"),
    ])

    # 7. Submission Requirements
    doc.add_h2("7.0", "SUBMISSION REQUIREMENTS")
    doc.add_body(f"{name} is requested to complete and return the following items with this document, duly stamped and signed:")
    doc.add_h3("7.1", "Company Information")
    ci_items = ["Company Name (Arabic)", "Company Name (English)", "Commercial Registration No.",
        "CR Issue Date / Expiry", "VAT Certificate No.", "Saudi Contractors Authority No.",
        "Classification Grade / Category", "ISO Certifications Held", "Year Established",
        "No. of Full-Time Employees", "Annual Turnover (SAR) — Last 3 Years", "Bank Reference"]
    t = doc.doc.add_table(rows=1+len(ci_items), cols=2)
    format_header_cell(t.cell(0, 0), "Item")
    format_header_cell(t.cell(0, 1), "Response")
    for i, item in enumerate(ci_items):
        format_data_cell(t.cell(i+1, 0), item, WD_ALIGN_PARAGRAPH.LEFT)
        format_data_cell(t.cell(i+1, 1), "___________________________", WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_body("")

    doc.add_h3("7.2", "Portfolio — Museum / Cultural Projects")
    doc.add_body("Please list all museum, cultural, or high-profile public space projects completed in the last 5 years. Attach project sheets with photos and client references.")
    t = doc.doc.add_table(rows=6, cols=6)
    for i, h in enumerate(["#", "Project Name", "Location", "Client", "Value (SAR)", "Year"]):
        format_header_cell(t.cell(0, i), h)
    for r in range(1, 6):
        for c in range(6):
            format_data_cell(t.cell(r, c), "", WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_body("")

    doc.add_h3("7.3", "Key Personnel — CVs Required")
    doc.add_body("Attach CVs for the following proposed team members:")
    t = doc.doc.add_table(rows=7, cols=3)
    format_header_cell(t.cell(0, 0), "Role")
    format_header_cell(t.cell(0, 1), "Min. Experience")
    format_header_cell(t.cell(0, 2), "Name of Proposed Person")
    for i, (role, exp) in enumerate([
        ("Lead Acoustician", "10+ years, museum/cultural project experience"),
        ("Acoustic Engineer", "7+ years"),
        ("Acoustic Modeller", "5+ years, CadnaR/Treble experience"),
        ("Site Acoustician", "5+ years, commissioning experience"),
        ("BIM Coordinator", "3+ years, Revit/Navisworks"),
        ("Project Manager", "5+ years"),
    ]):
        format_data_cell(t.cell(i+1, 0), role, WD_ALIGN_PARAGRAPH.LEFT)
        format_data_cell(t.cell(i+1, 1), exp, WD_ALIGN_PARAGRAPH.LEFT)
        format_data_cell(t.cell(i+1, 2), "___________________", WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_body("")

    doc.add_h3("7.4", "Method Statement")
    doc.add_body("Provide a method statement covering: approach to acoustic assessment, design methodology, treatment specification process, commissioning procedures, and coordination with MEP/AV trades.")
    doc.add_h3("7.5", "Programme")
    doc.add_body("Provide a milestone-based programme aligned with the project master schedule:")
    make_2col_table(doc, ["Milestone", "Target (Day-N)"], [
        ("Prequalification", "0"),
        ("Acoustic Assessment Report (post-survey)", "30"),
        ("50% Design Recommendations", "50"),
        ("90% Design Recommendations + RDS update", "65"),
        ("Treatment specifications + samples", "80"),
        ("Installation witnessing (continuous)", "150-210"),
    ])

    # 8. Coordination Interfaces
    doc.add_h2("8.0", "COORDINATION INTERFACES")
    make_3col_table(doc, ["Interface", "Counterpart", "Resolution"], [
        ("Acoustic ceiling vs HVAC ducts, sprinklers, lighting", "Samaya BIM Mgr (Dr. Waleed Salah)", "RCP federation @ LOD 300"),
        ("AV speaker coverage vs Acoustic", "NRS, AV Sub (Rawasin)", "Speaker coverage analysis with Acoustic"),
        ("HVAC equipment noise vs Gallery NC", "Samaya MEP Coord", "TAB, sound attenuation"),
        ("Speech privacy vs Setworks", "NRS (Eng. Jim)", "Setwork acoustic verification"),
        ("Baffle suspension vs Structural anchors", "Structural / Rigging Contractor", "M6 rod fixing schedule + load calculations"),
        ("PAVA speakers vs acoustic ceilings", "FLS Specialist", "Mounting plates or relocation"),
    ])

    # 9. Standards
    doc.add_h2("9.0", "STANDARDS AND REFERENCES")
    make_2col_table(doc, ["Standard / Reference", "Application"], [
        ("BS 8233:2014", "Sound insulation and noise reduction for buildings"),
        ("ISO 3382-1:2009", "Measurement of room acoustic parameters — reverberation time"),
        ("ISO 354 / ASTM C423", "Sound absorption testing"),
        ("ISO 717", "Sound insulation rating"),
        ("IEC 60268-16:2020", "Speech intelligibility (STI)"),
        ("DIN 18041", "Acoustic quality in rooms"),
        ("EN 13501-1", "Fire classification of construction products"),
        ("ASTM E84", "Surface burning characteristics"),
        ("ASHRAE", "HVAC noise criteria"),
        ("SBC 501", "Saudi Mechanical Code (noise)"),
        ("iAcoustics Strategy V2 (ref. J3574)", "Governing acoustic design — RT60/NR targets per space"),
        ("NRS Acoustic Materials Spec Rev 00", "Material specification — Tier A/B/C zones, NRC, fire, VOC"),
    ])

    # 10. Declaration
    doc.add_h2("10.0", "DECLARATION & COMPANY STAMP")
    doc.add_body(f"I, the undersigned, confirm that the information provided in this prequalification submission is complete and accurate. I confirm that {name} has the capability, resources, and experience to execute the Acoustic Specialist Consultant scope for the Aseer Regional Museum project as described in this document.")
    doc.add_body("")
    doc.add_body("__________________________________________________")
    doc.add_body("Authorised Signatory: ___________________________")
    doc.add_body("Signature:  ___________________________")
    doc.add_body("Date:  ___________________________")
    doc.add_body("Company Stamp:  ___________________________")
    doc.add_body("")
    doc.add_body("- End of Document -")

    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    for c in COMPANIES:
        path = generate(c)
        print(f"Generated: {path}")

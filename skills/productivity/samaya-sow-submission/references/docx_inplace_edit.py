"""
Reusable helpers for editing a Samaya/Sister SOW .docx IN-PLACE (preserves template,
header logo, table styles). Copy into a session and adapt — do NOT rebuild from empty Document().

Usage:
  d = docx.Document(SRC_TEMPLATE)
  set_para(find_para('Old heading'), 'New text')
  set_cell(table.rows[0].cells[1], 'New value')
  # insert after a paragraph:
  add_para_after(find_para('Anchor'), 'New line', bold=True)
  # delete a block between two markers:
  delete_between(d, '2.1  PHYSICAL COMPONENTS', '6.0  NRS RFI')
  # swap header logo:
  swap_header_logo(d, '/path/to/new-logo.jpg')
  d.save(OUT)
"""
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

def set_para(p, text):
    for r in list(p.runs): r._r.getparent().remove(r._r)
    p.add_run(text)

def set_cell(cell, text):
    for p in cell.paragraphs:
        for r in list(p.runs): r._r.getparent().remove(r._r)
    cell.paragraphs[0].add_run(text)

def add_para_after(ref, text, style=None, bold=False):
    new = docx.oxml.OxmlElement('w:p'); ref._p.addnext(new)
    p = Paragraph(new, ref._parent)
    if style: p.style = style
    r = p.add_run(text); r.bold = bold
    return p

def find_para(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith): return p
    return None

def delete_between(doc, start_marker, end_marker):
    """Remove all body children (paras + tables) strictly between two paragraph markers."""
    body = doc.element.body
    start = find_para(doc, start_marker)._p
    end = find_para(doc, end_marker)._p
    rm = []; collecting = False
    for child in list(body.iterchildren()):
        if child is start:
            collecting = True; continue
        if child is end:
            break
        if collecting: rm.append(child)
    for e in rm: e.getparent().remove(e)

def swap_header_logo(doc, new_logo_path):
    """Replace header image in all sections. Handles missing sections by adding sectPr."""
    # Ensure at least one section exists
    if len(doc.sections) == 0:
        from docx.oxml import OxmlElement
        sectPr = OxmlElement('w:sectPr')
        doc.element.body.append(sectPr)
    for sec in doc.sections:
        hdr = sec.header
        if hdr.is_linked_to_previous:
            hdr.is_linked_to_previous = False
        for rel in list(hdr.part.rels.values()):
            if 'image' in rel.reltype:
                with open(new_logo_path, 'rb') as f:
                    rel.target_part._blob = f.read()
                return True
    return False

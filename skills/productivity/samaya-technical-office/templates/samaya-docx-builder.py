"""
Samaya Document Builder v1.0
Usage: copy this file, import SamayaDocument, build your doc.
Full style guide: references/document-style-guide.md
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

class SamayaColors:
    NAVY = RGBColor(0x1E, 0x29, 0x3B)
    ACCENT_RED = RGBColor(0xB0, 0x1E, 0x2F)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    DARK_GRAY = RGBColor(0x33, 0x41, 0x55)
    MEDIUM_GRAY = RGBColor(0x64, 0x74, 0x8B)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    NAVY_HEX = '1E293B'
    LIGHT_GRAY_HEX = 'F1F5F9'
    WHITE_HEX = 'FFFFFF'
    BORDER_GRAY_HEX = 'CBD5E1'
    MUTED_BG_HEX = 'F8FAFC'

class SamayaDocument:
    def __init__(self, project_name='', doc_ref='', date=''):
        self.doc = Document()
        self._setup_page()
        if project_name:
            self._create_header(project_name, doc_ref, date)
        if doc_ref:
            self._create_footer(doc_ref)

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.2)

    def _create_header(self, project_name, doc_ref, date):
        header = self.doc.sections[0].header
        header.is_linked_to_previous = False
        for p in header.paragraphs: p.clear()
        p = header.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        parts = [
            (project_name, SamayaColors.DARK_GRAY, 7.5, False),
            ('  |  ', SamayaColors.MEDIUM_GRAY, 7.5, False),
            (doc_ref, SamayaColors.MEDIUM_GRAY, 7.5, False),
            ('  |  ', SamayaColors.MEDIUM_GRAY, 7.5, False),
            (date, SamayaColors.MEDIUM_GRAY, 7.5, False),
        ]
        for text, color, size, bold in parts:
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.font.bold = bold
        self._add_bottom_border(p, 'CBD5E1', 4)

    def _create_footer(self, doc_number):
        footer = self.doc.sections[0].footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs: p.clear()
        p = footer.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        self._add_top_border(p, '1E293B', 8)
        run = p.add_run(doc_number)
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = SamayaColors.MEDIUM_GRAY
        run = p.add_run('    ')
        run.font.size = Pt(8)
        run = p.add_run('Page ')
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = SamayaColors.MEDIUM_GRAY
        self._add_field(p, 'PAGE')
        run = p.add_run(' of ')
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = SamayaColors.MEDIUM_GRAY
        self._add_field(p, 'NUMPAGES')

    def add_title(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(8)
        pf.line_spacing = Pt(22)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run = p.add_run(text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = SamayaColors.NAVY
        self._add_bottom_border(p, '1E293B', 10)
        return p

    def add_h2(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(4)
        pf.line_spacing = Pt(16)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run = p.add_run(text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = SamayaColors.NAVY
        self._add_bottom_border(p, 'CBD5E1', 4)
        return p

    def add_body(self, text, size=10):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = Pt(12)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.color.rgb = SamayaColors.BLACK
        return p

    def add_table(self, headers, data, col_widths_cm=None):
        ncols = len(headers)
        table = self.doc.add_table(rows=1 + len(data), cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if col_widths_cm is None:
            col_widths_cm = [16.5 / ncols] * ncols
        for i, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(width)
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            self._set_cell_shading(cell, SamayaColors.NAVY_HEX)
            self._set_cell_border(cell)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(h.upper())
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = SamayaColors.WHITE
        for r_idx, row_data in enumerate(data):
            bg = SamayaColors.WHITE_HEX if r_idx % 2 == 0 else SamayaColors.LIGHT_GRAY_HEX
            for c_idx, val in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                self._set_cell_shading(cell, bg)
                self._set_cell_border(cell)
                run = cell.paragraphs[0].add_run(str(val))
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.color.rgb = SamayaColors.BLACK
                if r_idx == len(data) - 1:
                    run.font.bold = True
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(0)
        return table

    def add_note(self, text, label='NOTE'):
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = table.cell(0, 0)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
            f'  <w:left w:val="single" w:sz="18" w:space="4" w:color="B01E2F"/>'
            f'</w:tcBorders>')
        tcPr.append(tcBorders)
        self._set_cell_shading(cell, SamayaColors.MUTED_BG_HEX)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(f'{label}: ')
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = SamayaColors.ACCENT_RED
        run = cell.paragraphs[0].add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = SamayaColors.DARK_GRAY
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(3)
        sp.paragraph_format.space_after = Pt(0)
        return table

    def add_bullet(self, text, level=0):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.left_indent = Cm(0.5 + (level * 0.5))
        run = p.add_run(f'-  {text}')
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = SamayaColors.BLACK
        return p

    # -- Low-level helpers --
    def _set_cell_shading(self, cell, color_hex):
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)

    def _set_cell_border(self, cell, color='CBD5E1', size=4):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(size))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _add_bottom_border(self, paragraph, color='CBD5E1', size=4):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_top_border(self, paragraph, color='1E293B', size=8):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), str(size))
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), color)
        pBdr.append(top)
        pPr.append(pBdr)

    def _add_field(self, paragraph, field_code):
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)
        run2 = paragraph.add_run()
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = field_code
        run2._r.append(instr_text)
        run3 = paragraph.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld_char_end)

    def save(self, path):
        self.doc.save(path)

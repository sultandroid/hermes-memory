"""
Samaya Investment Company — Technical Office DOCX Template
==========================================================
Implements Samaya_Doc_Style_Guide_v1.0 for python-docx.

Usage:
    from samaya_doc_template import SamayaDoc, SamayaColors
    doc = SamayaDoc()
    doc.create_header('Aseer Museum', 'ASR-SAM-XXX-001', 'RPT', 'A', 'Jun 2026')
    doc.create_footer('ASR-SAM-XXX-001')
    doc.add_h1('DOCUMENT TITLE')
    doc.add_body('Your text here.')
    doc.add_table(headers, rows, col_widths_cm=[...])
    doc.save('output.docx')
"""

import datetime, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import docx.oxml

# ─── Color Palette ────────────────────────────────────────────
class SamayaColors:
    NAVY        = RGBColor(0x1E, 0x29, 0x3B)
    ACCENT_RED  = RGBColor(0xB0, 0x1E, 0x2F)
    WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
    DARK_GRAY   = RGBColor(0x33, 0x41, 0x55)
    MEDIUM_GRAY = RGBColor(0x64, 0x74, 0x8B)
    LIGHT_GRAY  = RGBColor(0xF1, 0xF5, 0xF9)
    BORDER_GRAY = RGBColor(0xCB, 0xD5, 0xE1)
    MUTED_BG    = RGBColor(0xF8, 0xFA, 0xFC)
    BLACK       = RGBColor(0x00, 0x00, 0x00)

    @staticmethod
    def hex_to_rgb(hex_color):
        h = hex_color.lstrip('#')
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

# ─── Helpers ──────────────────────────────────────────────────
def _font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color

def _border(paragraph, where='bottom', color='CBD5E1', size=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el = OxmlElement(f'w:{where}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(size))
    el.set(qn('w:space'), '1')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def _add_field(paragraph, code):
    r = paragraph.add_run(); fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), 'begin'); r._r.append(fc)
    r2 = paragraph.add_run(); it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve'); it.text = code; r2._r.append(it)
    r3 = paragraph.add_run(); fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end'); r3._r.append(fc2)

def _tabs(p, c=5670, r=11340):
    pPr = p._p.get_or_add_pPr(); tabs = OxmlElement('w:tabs')
    for v, pos in [('center', c), ('right', r)]:
        t = OxmlElement('w:tab'); t.set(qn('w:val'), v); t.set(qn('w:pos'), str(pos)); tabs.append(t)
    pPr.append(tabs)

def _shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'))

def _cmargins(cell, t=28, b=28, l=56, r_=56):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{t}" w:type="dxa"/><w:bottom w:w="{b}" w:type="dxa"/><w:start w:w="{l}" w:type="dxa"/><w:end w:w="{r_}" w:type="dxa"/></w:tcMar>'))

def _cborder(cell, color='CBD5E1', sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/><w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/></w:tcBorders>'))

def _hcell(cell, text):
    _shade(cell, '1E293B'); _cmargins(cell); _cborder(cell)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(1); pf.space_after = Pt(1)
    pf.line_spacing = Pt(11); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    for r in p.runs: r.clear()
    _font(p.add_run(text.upper()), size=9.5, bold=True, color=SamayaColors.WHITE)

def _dcell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    _cmargins(cell); _cborder(cell)
    p = cell.paragraphs[0]; p.alignment = align
    pf = p.paragraph_format; pf.space_before = Pt(1); pf.space_after = Pt(1)
    pf.line_spacing = Pt(11); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    for r in p.runs: r.clear()
    _font(p.add_run(str(text)), size=9.5, color=SamayaColors.BLACK)

# ─── Logo path ────────────────────────────────────────────────
LOGO = '/Users/mohamedessa/Library/CloudStorage/OneDrive-SAMAYAINVESTMENT/Samaya/Technical Office/_Style-Guides/logos archives/samaya-logo-trans.png'

# ─── Main Class ───────────────────────────────────────────────
class SamayaDoc:
    def __init__(self):
        self.doc = Document()
        s = self.doc.sections[0]
        s.page_width = Cm(21.0); s.page_height = Cm(29.7)
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
        s.header_distance = Cm(1.5); s.footer_distance = Cm(1.2)
        st = self.doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(6)

    # ── Header ──
    def create_header(self, project_name, doc_ref, doc_type, revision, date=None):
        if date is None: date = datetime.date.today().strftime('%b %Y')
        h = self.doc.sections[0].header; h.is_linked_to_previous = False
        for p in h.paragraphs: p.clear()
        p = h.paragraphs[0]; p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        if os.path.exists(LOGO):
            r = p.add_run(); r.add_picture(LOGO, width=Cm(2.5)); _font(p.add_run('\t'), size=7.5)
        else:
            _font(p.add_run('[SAMAYA]  \t'), size=7.5, bold=True, color=SamayaColors.NAVY)
        _font(p.add_run(f'{project_name}\n'), size=7.5, color=SamayaColors.DARK_GRAY)
        _font(p.add_run(f'Doc Ref: {doc_ref}  |  {doc_type}'), size=7.5, color=SamayaColors.MEDIUM_GRAY)
        _font(p.add_run('\t'), size=7.5)
        _font(p.add_run(f'Rev: {revision}\n'), size=7.5, bold=True, color=SamayaColors.NAVY)
        _font(p.add_run(date), size=7.5, color=SamayaColors.MEDIUM_GRAY)
        _tabs(p); _border(p, color='CBD5E1', size=4)

    # ── Footer ──
    def create_footer(self, doc_number, confidential=True):
        f = self.doc.sections[0].footer; f.is_linked_to_previous = False
        for p in f.paragraphs: p.clear()
        p = f.paragraphs[0]; p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(0)
        _border(p, 'top', color='1E293B', size=8)
        _font(p.add_run(doc_number), size=8, color=SamayaColors.MEDIUM_GRAY)
        _font(p.add_run('\t'), size=8)
        _font(p.add_run('Page '), size=8, color=SamayaColors.MEDIUM_GRAY); _add_field(p, 'PAGE')
        _font(p.add_run(' of '), size=8, color=SamayaColors.MEDIUM_GRAY); _add_field(p, 'NUMPAGES')
        _font(p.add_run('\t'), size=8)
        _font(p.add_run('Samaya Investment Company'), size=8, color=SamayaColors.MEDIUM_GRAY)
        _tabs(p)
        if confidential:
            p2 = f.add_paragraph(); p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
            _font(p2.add_run('This document is the property of Samaya Investment Company. Unauthorised reproduction or distribution is prohibited.'), size=7, italic=True, color=SamayaColors.MEDIUM_GRAY)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Headings ──
    def add_h1(self, text):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format; pf.space_before = Pt(18); pf.space_after = Pt(12)
        pf.line_spacing = Pt(24); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.keep_with_next = True
        _font(p.add_run(text.upper()), size=18, bold=True, color=SamayaColors.NAVY)
        _border(p, color='1E293B', size=12); return p

    def add_h2(self, number, text):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format; pf.space_before = Pt(18); pf.space_after = Pt(6)
        pf.line_spacing = Pt(18); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.keep_with_next = True
        _font(p.add_run(f'{number}  {text.upper()}'), size=14, bold=True, color=SamayaColors.NAVY)
        _border(p, color='CBD5E1', size=6); return p

    def add_h2_u(self, text):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format; pf.space_before = Pt(18); pf.space_after = Pt(6)
        pf.line_spacing = Pt(18); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.keep_with_next = True
        _font(p.add_run(text.upper()), size=14, bold=True, color=SamayaColors.NAVY)
        _border(p, color='CBD5E1', size=6); return p

    def add_h3(self, number, text):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format; pf.space_before = Pt(12); pf.space_after = Pt(4)
        pf.line_spacing = Pt(16); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; pf.keep_with_next = True
        _font(p.add_run(f'{number}  {text.upper()}'), size=12, bold=True, color=SamayaColors.DARK_GRAY)
        _border(p, color='CBD5E1', size=4); return p

    # ── Body ──
    def add_body(self, text, size=11, bold=False, italic=False, color=None, align=None):
        p = self.doc.add_paragraph(); p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(6)
        pf.line_spacing = Pt(13); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        _font(p.add_run(text), size=size, bold=bold, italic=italic, color=color or SamayaColors.BLACK)
        return p

    def add_rich_body(self, segments):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(6)
        pf.line_spacing = Pt(13); pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        for s in segments:
            _font(p.add_run(s.get('text','')), size=s.get('size',11), bold=s.get('bold',False), italic=s.get('italic',False), color=s.get('color',SamayaColors.BLACK))
        return p

    # ── Table ──
    def add_table(self, headers, rows, col_widths_cm=None):
        t = self.doc.add_table(rows=1+len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
        if col_widths_cm:
            for row in t.rows:
                for i, w in enumerate(col_widths_cm): row.cells[i].width = Cm(w)
        for i, h in enumerate(headers): _hcell(t.rows[0].cells[i], h)
        for r, rd in enumerate(rows):
            bg = 'F1F5F9' if r % 2 == 1 else 'FFFFFF'
            for c, v in enumerate(rd):
                cl = t.rows[r+1].cells[c]
                _dcell(cl, v, align=WD_ALIGN_PARAGRAPH.CENTER if c==0 else WD_ALIGN_PARAGRAPH.LEFT)
                _shade(cl, bg)
        return t

    def line(self):
        p = self.doc.add_paragraph(); pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(2); pf.line_spacing = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY; return p

    def save(self, path): self.doc.save(path); return path

    def save_temp(self, prefix='samaya_doc_'):
        ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        path = os.path.join('/tmp', f'{prefix}{ts}.docx')
        self.doc.save(path); return path

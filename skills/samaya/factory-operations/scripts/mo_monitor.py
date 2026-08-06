#!/usr/bin/env python3
"""
نظام متابعة المخزون — مصنع سمايا
المصدر: Odoo (قراءة فقط + ملاحظات شاتر)
التخزين: samaya-profile repo
التنبيهات: تقارير + إشعارات

المهام:
1. كل يومين: تذكير بصرف المواد على MOs بدون مواد
2. أسبوعياً: تقرير Word كامل (مخزون + POs + MOs)
3. فحص: الرصيد السالب والمواد النافدة
"""
import os, sys, json, datetime, re
from collections import defaultdict
from docx import Document

# === Odoo Connection ===
import xmlrpc.client, ssl
ENV_PATH = os.path.expanduser('~/.config/samaya/odoo.env')
env = {}
with open(ENV_PATH) as f:
    for line in f:
        line = line.strip()
        if not line or line.startswith('#') or '=' not in line: continue
        k, v = line.split('=', 1)
        env[k.strip()] = v.strip()

url = env['ODOO_URL']
db = env['ODOO_DB']
user = env['ODOO_USER']
key = env['ODOO_API_KEY']
ctx = ssl.create_default_context()
common = xmlrpc.client.ServerProxy(f'{url}/xmlrpc/2/common', context=ctx)
uid = common.authenticate(db, user, key, {})
models = xmlrpc.client.ServerProxy(f'{url}/xmlrpc/2/object', context=ctx)

MODE = sys.argv[1] if len(sys.argv) > 1 else 'check'
REPO_DIR = '/Users/mohamedessa/samaya-profile/00_Admin'
REMINDER_MARKER = '__REMINDER_MARKER__'

# === HELPERS ===
def get_all_child_ids(cat_id, all_cats):
    ids = [cat_id]
    for c in all_cats:
        if c['parent_id'] and c['parent_id'][0] == cat_id:
            ids.extend(get_all_child_ids(c['id'], all_cats))
    return ids

def is_workday():
    now = datetime.datetime.now()
    wd = now.weekday()
    if wd == 4:  # Friday
        return False
    hour = now.hour
    if hour < 8 or hour >= 17:
        return False
    return True

def should_run_today():
    if not is_workday():
        return False
    days_since_epoch = (datetime.datetime.now().date() - datetime.date(2026, 1, 1)).days
    return days_since_epoch % 2 == 0

# === 1. MOs بدون مواد ===
def check_mos_without_materials():
    two_days_ago = (datetime.datetime.now() - datetime.timedelta(days=2)).strftime('%Y-%m-%d')
    mo_ids = models.execute_kw(db, uid, key, 'mrp.production', 'search', [
        [('state', 'in', ['progress', 'confirmed']), ('create_date', '<', two_days_ago)]
    ])
    if not mo_ids:
        return []
    mos = models.execute_kw(db, uid, key, 'mrp.production', 'read', [mo_ids], 
        {'fields': ['id', 'name', 'product_id', 'state', 'create_date', 'move_raw_ids', 'date_start']})
    return [mo for mo in mos if not mo.get('move_raw_ids')]

# === 2. MOs في to_close بدون مواد ===
def check_to_close_without_materials():
    to_close_ids = models.execute_kw(db, uid, key, 'mrp.production', 'search', [[('state', '=', 'to_close')]])
    if not to_close_ids:
        return []
    mos = models.execute_kw(db, uid, key, 'mrp.production', 'read', [to_close_ids], 
        {'fields': ['id', 'name', 'product_id', 'move_raw_ids', 'date_finished']})
    return [mo for mo in mos if not mo.get('move_raw_ids')]

# === 3. مسح التذكيرات القديمة ===
def delete_old_reminders(mo_id):
    msgs = models.execute_kw(db, uid, key, 'mail.message', 'search_read', [
        [('model', '=', 'mrp.production'), ('res_id', '=', mo_id)]
    ], {'fields': ['id', 'body', 'date', 'author_id'], 'limit': 50})
    deleted = 0
    for m in msgs:
        body = m.get('body', '') or ''
        if REMINDER_MARKER in body:
            try:
                models.execute_kw(db, uid, key, 'mail.message', 'unlink', [[m['id']]])
                deleted += 1
            except:
                pass
    return deleted

# === 4. وضع تذكير جديد ===
def post_reminder(mo):
    mo_id = mo['id']
    mo_name = mo['name']
    prod = mo['product_id'][1] if mo['product_id'] else '?'
    qty = mo.get('product_qty', 0)
    date_start = (mo.get('date_start') or '')[:10]
    delete_old_reminders(mo_id)
    body = f'''<p>🔔 <strong>تذكير بصرف المواد</strong> <span style="display:none">{REMINDER_MARKER}</span></p><p>يرجى تسجيل المواد الخاصة بأمر التصنيع <strong>{mo_name}</strong></p><p>المنتج: {prod}</p><p>الكمية: {qty}</p><p>تاريخ البدء: {date_start}</p><p>ملاحظة: لا يمكن إغلاق أمر التصنيع بدون تسجيل المواد ووقت مراكز العمل</p>'''
    try:
        models.execute_kw(db, uid, key, 'mrp.production', 'message_post', [mo_id], {
            'body': body, 'message_type': 'comment', 'subtype_xmlid': 'mail.mt_comment'
        })
        return True
    except Exception as e:
        print(f'  ERROR: {e}')
        return False

# === 5. المخزون ===
def get_inventory_summary():
    factory_locs = [45, 46, 47, 51, 77, 78]
    raw_cat_ids = models.execute_kw(db, uid, key, 'product.category', 'search', [
        [('name', 'in', ['Raw materials', 'Raw Materials'])]
    ])
    all_cats = models.execute_kw(db, uid, key, 'product.category', 'search_read', [[]], 
        {'fields': ['id', 'name', 'parent_id']})
    raw_cat_tree = list(set(sum((get_all_child_ids(rc, all_cats) for rc in raw_cat_ids), [])))
    
    products = models.execute_kw(db, uid, key, 'product.product', 'search_read', [
        [('categ_id', 'in', raw_cat_tree), ('tracking', '!=', 'none')]
    ], {'fields': ['id', 'display_name', 'default_code', 'qty_available', 'categ_id']})
    
    quants = models.execute_kw(db, uid, key, 'stock.quant', 'search_read', [
        [('product_id', 'in', [p['id'] for p in products]), ('location_id', 'in', factory_locs)]
    ], {'fields': ['product_id', 'location_id', 'quantity', 'reserved_quantity']})
    
    quant_by_prod = defaultdict(float)
    reserved_by_prod = defaultdict(float)
    for q in quants:
        pid = q['product_id'][0]
        quant_by_prod[pid] += q['quantity'] or 0
        reserved_by_prod[pid] += q['reserved_quantity'] or 0
    
    date_from = (datetime.datetime.now() - datetime.timedelta(days=365)).strftime('%Y-%m-%d')
    moves = models.execute_kw(db, uid, key, 'stock.move', 'search_read', [
        [('product_id', 'in', [p['id'] for p in products]), ('date', '>=', date_from),
         ('state', '=', 'done')]
    ], {'fields': ['product_id', 'product_qty', 'location_id', 'location_dest_id', 'date'],
        'limit': 10000})
    
    incoming = defaultdict(float)
    outgoing = defaultdict(float)
    for m in moves:
        pid = m['product_id'][0]
        qty = m['product_qty']
        loc_to = m['location_dest_id'][0] if m['location_dest_id'] else 0
        loc_from = m['location_id'][0] if m['location_id'] else 0
        if loc_to in factory_locs:
            incoming[pid] += qty
        elif loc_from in factory_locs and loc_to in [15, 98]:
            outgoing[pid] += qty
    
    summary = []
    for p in products:
        pid = p['id']
        on_hand = quant_by_prod.get(pid, 0)
        available = on_hand - reserved_by_prod.get(pid, 0)
        out = outgoing.get(pid, 0)
        summary.append({
            'code': p.get('default_code', '') or '',
            'name': p['display_name'],
            'category': p['categ_id'][1] if p['categ_id'] else '',
            'on_hand': on_hand,
            'available': available,
            'incoming_12m': incoming.get(pid, 0),
            'outgoing_12m': out,
            'monthly_avg': round(out / 12, 1),
            'coverage_days': round((on_hand / (out / 365)) if out > 0 else 999, 1)
        })
    return summary

# === 6. التقرير الأسبوعي ===
def generate_report():
    now = datetime.datetime.now()
    week_ago = (now - datetime.timedelta(days=7)).strftime('%Y-%m-%d')
    
    new_mos = models.execute_kw(db, uid, key, 'mrp.production', 'search_read', [
        [('create_date', '>=', week_ago)]
    ], {'fields': ['name', 'product_id', 'state', 'create_date', 'move_raw_ids', 'bom_id'],
        'order': 'create_date DESC'})
    done_mos = models.execute_kw(db, uid, key, 'mrp.production', 'search_read', [
        [('date_finished', '>=', week_ago), ('state', '=', 'done')]
    ], {'fields': ['name', 'product_id', 'state', 'date_finished', 'move_raw_ids'],
        'order': 'date_finished DESC'})
    
    no_mat_mos = [m for m in new_mos if not m.get('move_raw_ids')]
    no_bom_mos = [m for m in new_mos if not m.get('bom_id')]
    inventory = get_inventory_summary()
    neg_stock = [i for i in inventory if i['on_hand'] < 0]
    low_stock = [i for i in inventory if 0 < i['on_hand'] < 5 and i['outgoing_12m'] > 0]
    zero_stock = [i for i in inventory if i['on_hand'] == 0 and i['outgoing_12m'] > 0]
    
    pending_pos = models.execute_kw(db, uid, key, 'purchase.order', 'search_read', [
        [('project_id', '=', 244), ('state', 'in', ['purchase', 'sent'])]
    ], {'fields': ['name', 'partner_id', 'amount_total', 'date_order', 'receipt_status'],
        'order': 'date_order DESC'})
    no_receipt = [p for p in pending_pos if p.get('receipt_status') in (False, 'False', 'pending')]
    
    # Build DOCX
    doc = Document()
    title = doc.add_heading('تقرير المصنع الأسبوعي', level=0)
    title.alignment = 1
    doc.add_paragraph(f'الفترة: {week_ago} → {now.strftime("%Y-%m-%d")}')
    doc.add_paragraph(f'تاريخ التقرير: {now.strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('')
    
    doc.add_heading('1. أوامر التصنيع', level=1)
    doc.add_heading(f'جديد: {len(new_mos)}', level=2)
    t = doc.add_table(rows=1, cols=4); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(['رقم الأمر', 'المنتج', 'الحالة', 'مواد مسجلة']):
        t.rows[0].cells[i].text = h
    for m in new_mos[:20]:
        r = t.add_row().cells
        r[0].text = m['name']
        r[1].text = m['product_id'][1] if m['product_id'] else '?'
        r[2].text = m['state']
        r[3].text = 'نعم' if m.get('move_raw_ids') else 'لا'
    
    doc.add_paragraph('')
    doc.add_heading(f'منجز: {len(done_mos)}', level=2)
    t2 = doc.add_table(rows=1, cols=4); t2.style = 'Light Grid Accent 1'
    for i, h in enumerate(['رقم الأمر', 'المنتج', 'تاريخ الإنجاز', 'مواد مسجلة']):
        t2.rows[0].cells[i].text = h
    for m in done_mos[:20]:
        r = t2.add_row().cells
        r[0].text = m['name']
        r[1].text = m['product_id'][1] if m['product_id'] else '?'
        r[2].text = (m.get('date_finished') or '')[:10]
        r[3].text = 'نعم' if m.get('move_raw_ids') else 'لا'
    
    doc.add_paragraph('')
    doc.add_heading('2. مخالفات التسجيل', level=1)
    doc.add_paragraph(f'أوامر بدون مواد مسجلة: {len(no_mat_mos)}')
    for m in no_mat_mos[:10]:
        doc.add_paragraph(f'  • {m["name"]} — {m["product_id"][1] if m["product_id"] else "?"}', style='List Bullet')
    doc.add_paragraph(f'أوامر بدون BoM: {len(no_bom_mos)}')
    for m in no_bom_mos[:10]:
        doc.add_paragraph(f'  • {m["name"]} — {m["product_id"][1] if m["product_id"] else "?"}', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('3. المخزون', level=1)
    doc.add_paragraph(f'رصيد سالب: {len(neg_stock)}')
    for i in neg_stock[:10]:
        doc.add_paragraph(f'  • {i["code"]} — {i["name"][:50]} — رصيد: {i["on_hand"]:.1f}', style='List Bullet')
    doc.add_paragraph(f'رصيد منخفض (<5): {len(low_stock)}')
    for i in low_stock[:10]:
        doc.add_paragraph(f'  • {i["code"]} — {i["name"][:50]} — رصيد: {i["on_hand"]:.1f}', style='List Bullet')
    doc.add_paragraph(f'نافد (صفر مع صرف): {len(zero_stock)}')
    for i in zero_stock[:10]:
        doc.add_paragraph(f'  • {i["code"]} — {i["name"][:50]} — صرف شهري: {i["monthly_avg"]:.1f}', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('4. أوامر الشراء المعلقة', level=1)
    doc.add_paragraph(f'POs بدون استلام كامل: {len(no_receipt)}')
    t3 = doc.add_table(rows=1, cols=4); t3.style = 'Light Grid Accent 1'
    for i, h in enumerate(['رقم PO', 'المورد', 'المبلغ', 'الاستلام']):
        t3.rows[0].cells[i].text = h
    for p in no_receipt[:15]:
        r = t3.add_row().cells
        r[0].text = p['name']
        r[1].text = p['partner_id'][1][:30] if p['partner_id'] else '?'
        r[2].text = f"{p.get('amount_total', 0):,.0f} SAR"
        r[3].text = str(p.get('receipt_status', '?'))
    
    report_path = f'{REPO_DIR}/weekly_report_{now.strftime("%Y%m%d")}.docx'
    doc.save(report_path)
    
    json_data = {
        'generated': now.strftime('%Y-%m-%d %H:%M'),
        'new_mos': len(new_mos), 'done_mos': len(done_mos),
        'no_material_mos': len(no_mat_mos), 'no_bom_mos': len(no_bom_mos),
        'negative_stock': len(neg_stock), 'low_stock': len(low_stock),
        'zero_stock': len(zero_stock), 'pending_pos': len(no_receipt),
        'total_raw_materials': len(inventory),
        'total_on_hand': sum(i['on_hand'] for i in inventory)
    }
    with open(f'{REPO_DIR}/inventory_snapshot.json', 'w') as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)
    
    return report_path, json_data

# === MAIN ===
if MODE == 'remind':
    if not should_run_today():
        print('اليوم ليس يوم تذكير')
        sys.exit(0)
    print('=== تذكير — MOs بدون مواد ===')
    mos = check_mos_without_materials()
    if not mos:
        print('✅ كل MOs مسجل فيها المواد')
    else:
        print(f'⚠️ {len(mos)} MO بدون مواد:')
        for mo in mos:
            prod = mo['product_id'][1] if mo['product_id'] else '?'
            print(f'  {mo["name"]} | {prod} | بدأ: {(mo.get("date_start") or "")[:10]}')
            ok = post_reminder(mo)
            print(f'    → {"✅" if ok else "❌"}')
    blocked = check_to_close_without_materials()
    if blocked:
        print(f'\n🔴 {len(blocked)} MO في to_close بدون مواد:')
        for mo in blocked:
            prod = mo['product_id'][1] if mo['product_id'] else '?'
            print(f'  {mo["name"]} | {prod}')

elif MODE == 'weekly':
    print('=== التقرير الأسبوعي ===')
    report_path, data = generate_report()
    print(f'✅ Word: {report_path}')
    print(f'✅ JSON: {REPO_DIR}/inventory_snapshot.json')
    print(f'\nملخص: MOs جديدة={data["new_mos"]}, منجزة={data["done_mos"]}, بدون مواد={data["no_material_mos"]}, بدون BoM={data["no_bom_mos"]}, رصيد سالب={data["negative_stock"]}, منخفض={data["low_stock"]}, نافد={data["zero_stock"]}, POs معلقة={data["pending_pos"]}')

elif MODE == 'check':
    print('=== فحص سريع ===')
    mos = check_mos_without_materials()
    print(f'MOs بدون مواد: {len(mos)}')
    for mo in mos[:5]:
        prod = mo['product_id'][1] if mo['product_id'] else '?'
        print(f'  {mo["name"]} | {prod}')
    blocked = check_to_close_without_materials()
    print(f'MOs في to_close بدون مواد: {len(blocked)}')
    inventory = get_inventory_summary()
    neg = [i for i in inventory if i['on_hand'] < 0]
    low = [i for i in inventory if 0 < i['on_hand'] < 5 and i['outgoing_12m'] > 0]
    zero = [i for i in inventory if i['on_hand'] == 0 and i['outgoing_12m'] > 0]
    print(f'رصيد سالب: {len(neg)}, منخفض: {len(low)}, نافد: {len(zero)}')
    print('✅ تم')
